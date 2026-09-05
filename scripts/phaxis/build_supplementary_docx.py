#!/usr/bin/env python3
"""Build the reviewer-visible PHAxis anonymized supplementary DOCX.

The builder shares the main-manuscript ``narrative_proposal`` typography and
the named ``plant_phenomics_manuscript`` override.  ``layout-fixture`` mode is
visibly provisional and creates nine placeholder figure plates.  ``final``
mode accepts a hash-closed final main manuscript/compile receipt and a
nine-figure supplementary suite bound to the same PHAxis 1.0.0 model
identities.  It never consumes author/title-page metadata.

This command is CPU-only.  It does not discover evidence, read annotations, or
start a GPU program.
"""

from __future__ import annotations

import argparse
from datetime import datetime, timezone
import hashlib
import importlib.util
import json
import os
from pathlib import Path
import re
import sys
import tempfile
from typing import Any, Mapping, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image


def _load_main_builder():
    path = Path(__file__).with_name("build_submission_docx.py")
    spec = importlib.util.spec_from_file_location("phaxis_main_docx_support", path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"cannot load DOCX support module: {path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


main = _load_main_builder()

from phaxis.supplementary_tables import (  # noqa: E402
    SupplementaryTableError,
    validate_supplementary_table_data_bundle,
)

SCHEMA_VERSION = "PHAxis-supplementary-docx-build-2.0"
FIGURE_SUITE_SCHEMA = main.FIGURE_SUITE_SCHEMA
FINAL_FIGURE_STATUS = main.FINAL_FIGURE_STATUS
PLACEHOLDER = re.compile(r"\{\{[A-Z0-9_]+\}\}")
HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
NUMBERED_ITEM = re.compile(r"^(\d+)\.\s+(.+?)\s*$")
INTERNAL_SECTION = "## Supplementary machine-fill policy"
SUPPLEMENTARY_FIGURES = (
    (1, "Supplementary_Figure_S01_stageb_input_architecture_targets", "Stage-B physical input representation, multihead architecture, and target contract"),
    (2, "Supplementary_Figure_S02_split_selection_development_strata", "Family-isolated split, operating-point selection, and development strata"),
    (3, "Supplementary_Figure_S03_identity_attachment_endpoint_assurance", "Identity, formal attachment, endpoint, and conditional-length assurance"),
    (4, "Supplementary_Figure_S04_primary_root_trait_agreement", "Agreement of 19 derived primary-root descriptors"),
    (5, "Supplementary_Figure_S05_provider_tiling_numerical_equivalence", "Root-provider equivalence, same-component root continuity, formal attachment, and tiled-inference assurance"),
    (6, "Supplementary_Figure_S06_expanded_overlay_gallery", "Expanded acquisition-challenge overlay gallery"),
    (7, "Supplementary_Figure_S07_biological_sensitivity_observability", "Clean-cohort D15 analysis, full-cohort D15 sensitivity, and observability"),
    (8, "Supplementary_Figure_S08_runtime_memory_io", "Direct runtime, memory, utilization, and I/O decomposition"),
    (9, "Supplementary_Figure_S09_multitrait_atlas_coverage_effect_heatmap", "Clean-cohort D15 four-condition phenotype map and coverage of all 32 descriptors"),
)


class SupplementaryDocxError(RuntimeError):
    """A supplementary source, evidence, or layout contract is invalid."""


def _require(condition: bool, message: str) -> None:
    if not condition:
        raise SupplementaryDocxError(message)


def _read_text(path: Path, role: str) -> tuple[bytes, str]:
    _require(path.is_file() and not path.is_symlink(), f"{role} does not exist or is a symlink: {path}")
    raw = path.read_bytes()
    try:
        return raw, raw.decode("utf-8")
    except UnicodeError as error:
        raise SupplementaryDocxError(f"{role} must be UTF-8") from error


def _read_json(path: Path, role: str) -> tuple[bytes, dict[str, Any]]:
    try:
        return main._read_json(path, role)
    except main.SubmissionDocxError as error:
        raise SupplementaryDocxError(str(error)) from error


def _extract_frontmatter(text: str) -> dict[str, str]:
    lines = text.splitlines()
    _require(lines and lines[0].startswith("# "), "supplement has no H1 title")
    result = {"title": lines[0][2:].strip()}
    labels = {
        "Software and model-system version": "software_version",
        "Companion main manuscript": "companion_main_manuscript",
        "Status": "status",
    }
    for line in lines[1:12]:
        match = re.match(r"^\*\*(.+?):\*\*\s*(.*?)\s*$", line)
        if match and match.group(1) in labels:
            result[labels[match.group(1)]] = match.group(2).strip(" `")
    for field in ("title", "software_version", "companion_main_manuscript", "status"):
        _require(bool(result.get(field)), f"supplement frontmatter lacks {field}")
    _require(result["software_version"] == "PHAxis 1.0.0", "supplement software identity changed")
    return result


def _validate_final_inputs(
    *,
    supplement_text: str,
    main_text: str,
    main_raw: bytes,
    compile_raw: bytes,
    compile_payload: Mapping[str, Any],
) -> None:
    upper = supplement_text.upper()
    _require(PLACEHOLDER.search(supplement_text) is None, "final supplement retains machine tokens")
    for marker in ("EVIDENCE-BOUND MASTER", "FINAL RENDERING REQUIRES", "PROVISIONAL", "LAYOUT FIXTURE"):
        _require(marker not in upper, f"final supplement retains provisional marker: {marker}")
    try:
        main._validate_compile_receipt(
            compile_payload,
            receipt_raw=compile_raw,
            manuscript_raw=main_raw,
            manuscript_text=main_text,
        )
    except main.SubmissionDocxError as error:
        raise SupplementaryDocxError(str(error)) from error
    _require(compile_payload.get("blind_images_used") == 0, "main compile receipt used blind images")
    _require(
        compile_payload.get("root_cap_region_statistics_included") is False,
        "main compile receipt includes root-cap region statistics",
    )
    return None


def _validate_figure_suite(
    payload: Mapping[str, Any],
    *,
    summary_path: Path,
    compile_payload: Mapping[str, Any],
) -> tuple[list[tuple[int, str, Path]], dict[str, Any]]:
    try:
        # Reuse the journal figure summary already validated for the six main
        # plates.  This keeps main and supplementary images on one model,
        # figure-input, train399, and blind-use identity graph.
        main._validate_figure_summary(
            payload,
            summary_path=summary_path,
            compile_receipt=compile_payload,
        )
    except main.SubmissionDocxError as error:
        raise SupplementaryDocxError(str(error)) from error
    claim = payload.get("claim_contract")
    _require(
        isinstance(claim, Mapping)
        and claim.get("supplementary_figure_count") == 9
        and claim.get("supplementary_table_data_resource_count") == 10,
        "publication figure suite must declare nine figures and ten table/data resources",
    )
    receipt_relative = payload.get("supplementary_table_bundle_receipt")
    _require(
        isinstance(receipt_relative, str) and bool(receipt_relative),
        "supplementary table/data receipt path is missing",
    )
    receipt_path = (summary_path.parent / receipt_relative).resolve()
    _require(
        receipt_path.is_relative_to(summary_path.parent.resolve()),
        "supplementary table/data receipt escapes the figure suite",
    )
    try:
        table_bundle = validate_supplementary_table_data_bundle(
            receipt_path, require_final=True
        )
    except SupplementaryTableError as error:
        raise SupplementaryDocxError(
            f"supplementary Table/Data S1--S10 validation failed: {error}"
        ) from error
    _require(
        payload.get("supplementary_tables") == table_bundle["items"]
        and payload.get("supplementary_table_bundle_receipt_sha256")
        == table_bundle["receipt_sha256"]
        and payload.get("supplementary_table_bundle_identity_sha256")
        == table_bundle["bundle_identity_sha256"]
        and payload.get("supplementary_table_bundle_sha256")
        == table_bundle["bundle_file_sha256"],
        "supplementary table/data summary binding differs",
    )
    figures = payload.get("supplementary_figures")
    bundle_hashes = payload.get("supplementary_figure_bundle_sha256")
    identity = payload.get("supplementary_figure_bundle_identity_sha256")
    expected_order = [stem for _, stem, _ in SUPPLEMENTARY_FIGURES]
    _require(
        isinstance(figures, Mapping) and list(figures) == expected_order,
        "supplementary figure order/set is invalid",
    )
    _require(
        isinstance(bundle_hashes, Mapping) and list(bundle_hashes) == expected_order,
        "supplementary figure hash order/set is invalid",
    )
    _require(
        identity == main._canonical_hash(bundle_hashes),
        "supplementary figure-bundle identity mismatch",
    )
    base = summary_path.parent.resolve()
    expected = {stem: (number, title) for number, stem, title in SUPPLEMENTARY_FIGURES}
    result: list[tuple[int, str, Path]] = []
    for stem, record in figures.items():
        _require(isinstance(record, dict), "supplementary figure record must be an object")
        number, default_title = expected[stem]
        _require(record.get("number") == f"S{number}", f"supplementary figure {stem} number changed")
        _require(record.get("status") == "final", f"supplementary figure {number} is not final")
        bundle = record.get("bundle")
        hashes = bundle_hashes[stem]
        _require(isinstance(bundle, Mapping) and isinstance(hashes, Mapping), f"supplementary figure {number} bundle is invalid")
        files = bundle.get("files")
        _require(isinstance(files, Mapping) and isinstance(files.get("png"), str), f"supplementary figure {number} PNG is missing")
        try:
            image_path = main._child_file(Path(files["png"]), base, f"supplementary figure {number} PNG")
        except main.SubmissionDocxError as error:
            raise SupplementaryDocxError(str(error)) from error
        _require(hashes.get("png") == main._sha256_file(image_path), f"supplementary figure {number} PNG hash mismatch")
        with Image.open(image_path) as image:
            width, height = image.size
        _require(width >= 1000 and height >= 600, f"supplementary figure {number} PNG is too small")
        title = record.get("title", default_title)
        _require(isinstance(title, str) and title.strip(), f"supplementary figure {number} title is empty")
        result.append((number, title.strip(), image_path))
    result.sort(key=lambda item: item[0])
    return result, table_bundle


def _add_table_data_index(document: Document, table_bundle: Mapping[str, Any]) -> None:
    """Add a compact human-readable index to the machine-readable S1--S10 files."""

    heading = document.add_paragraph()
    heading.style = "Heading 1"
    heading.add_run("Supplementary Table/Data File Index")
    rows: list[list[str]] = [
        [
            "Resource",
            "Title",
            "Bundle-relative files",
            "Denominator closure",
            "Item identity (SHA-256)",
        ],
    ]
    for record in table_bundle["items"].values():
        denominator = record["denominator_contract"]
        relative_files = [str(record["item_receipt"])]
        relative_files.extend(
            (Path(str(record["directory"])) / relative).as_posix()
            for relative in record["file_sha256"]
        )
        rows.append(
            [
                str(record["number"]),
                str(record["title"]),
                "\n".join(relative_files),
                str(denominator["closure_status"]),
                str(record["item_identity_sha256"]),
            ]
        )
    main._add_table(document, rows, landscape=False)


def _clean_body_lines(text: str) -> list[str]:
    lines = text.splitlines()
    _require(lines and lines[0].startswith("# "), "supplement title is missing")
    cleaned: list[str] = []
    for line in lines[1:]:
        # The machine-fill policy is a source-side validation contract.  It is
        # intentionally retained in the auditable Markdown master but is not
        # reader-facing supplementary material.
        if line.strip() == INTERNAL_SECTION:
            break
        if re.match(r"^\*\*(?:Software and model-system version|Companion main manuscript|Status):\*\*", line):
            continue
        cleaned.append(line)
    return cleaned


def _blocks(lines: Sequence[str]) -> list[tuple[str, Any]]:
    blocks: list[tuple[str, Any]] = []
    paragraph: list[str] = []

    def flush() -> None:
        if paragraph:
            blocks.append(("paragraph", " ".join(value.strip() for value in paragraph)))
            paragraph.clear()

    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line.strip():
            flush()
            index += 1
            continue
        heading = HEADING.match(line)
        if heading:
            flush()
            blocks.append(("heading", (len(heading.group(1)), heading.group(2))))
            index += 1
            continue
        numbered = NUMBERED_ITEM.match(line)
        if numbered:
            flush()
            blocks.append(("numbered", (int(numbered.group(1)), numbered.group(2))))
            index += 1
            continue
        if line.lstrip().startswith("|"):
            flush()
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                rows.append(main._split_table_row(lines[index]))
                index += 1
            _require(len(rows) >= 2, "supplementary Markdown table has no separator")
            _require(all(main.TABLE_SEPARATOR_CELL.fullmatch(cell) for cell in rows[1]), "supplementary Markdown table separator is invalid")
            blocks.append(("table", [rows[0], *rows[2:]]))
            continue
        paragraph.append(line)
        index += 1
    flush()
    return blocks


def _add_title_page(
    document: Document,
    *,
    frontmatter: Mapping[str, str],
    fixture: bool,
) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(
        "LAYOUT FIXTURE - NOT FOR SUBMISSION"
        if fixture
        else "Plant Phenomics | ANONYMIZED SUPPLEMENTARY MATERIALS"
    )
    main._set_font(run, size=11, bold=True, color="A61B1B" if fixture else "555555")
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(frontmatter["title"])
    main._set_font(run, size=16, bold=True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(
        "ANONYMOUS REVIEW COPY"
        if not fixture
        else "ANONYMOUS LAYOUT FIXTURE"
    )
    main._set_font(run, size=10, bold=True, color="A61B1B" if fixture else "555555")
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(8)
    for label, value in (
        ("Software", frontmatter["software_version"]),
        ("Supplementary figures", "9"),
        ("Supplementary tables/data resources", "10"),
        ("Companion main manuscript", frontmatter["companion_main_manuscript"]),
    ):
        row = document.add_paragraph()
        row.paragraph_format.left_indent = Inches(0.7)
        row.paragraph_format.right_indent = Inches(0.7)
        row.paragraph_format.space_after = Pt(3)
        run = row.add_run(f"{label}: ")
        main._set_font(run, size=10, bold=True)
        run = row.add_run(value)
        main._set_font(run, size=10)


def _add_body(
    document: Document,
    *,
    text: str,
    running_title: str,
    fixture: bool,
    number_id: int,
) -> int:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    main._set_page_geometry(section, landscape=False, line_numbers=True)
    main._configure_header_footer(section, running_title=running_title, fixture=fixture)
    table_count = 0
    prior_list_number = 0
    for kind, content in _blocks(_clean_body_lines(text)):
        if kind == "heading":
            level, title = content
            if level == 2 and title in {"Supplementary Figure Legends", "Supplementary Tables and Data Files"}:
                document.add_page_break()
            style = "Heading 1" if level == 2 else "Heading 2" if level == 3 else "Heading 3"
            paragraph = document.add_paragraph(style=style)
            main._add_inline(paragraph, title, base_size=14 if style == "Heading 1" else 12 if style == "Heading 2" else 11)
            main._set_keep_with_next(paragraph)
            prior_list_number = 0
        elif kind == "numbered":
            number, value = content
            _require(number == prior_list_number + 1, "supplementary numbered list is not contiguous")
            prior_list_number = number
            paragraph = document.add_paragraph()
            main._apply_numbering(paragraph, number_id)
            main._add_inline(paragraph, value, base_size=11)
        elif kind == "table":
            table_count += 1
            main._add_table(document, content, landscape=False)
            prior_list_number = 0
        else:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.5
            main._add_inline(paragraph, content, base_size=11)
            prior_list_number = 0
    return table_count


def _add_figure_plates(
    document: Document,
    *,
    figures: Sequence[tuple[int, str, Path]],
    running_title: str,
    fixture: bool,
) -> None:
    records = list(figures) if figures else [(n, title, Path()) for n, _, title in SUPPLEMENTARY_FIGURES]
    section = document.add_section(WD_SECTION.NEW_PAGE)
    main._set_page_geometry(section, landscape=False, line_numbers=False)
    main._configure_header_footer(section, running_title=running_title, fixture=fixture)
    for index, (number, title, image_path) in enumerate(records):
        if index:
            document.add_page_break()
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(f"Figure S{number}. {title}")
        main._set_font(run, size=10, bold=True)
        if fixture:
            box = document.add_paragraph()
            box.alignment = WD_ALIGN_PARAGRAPH.CENTER
            box.paragraph_format.space_before = Pt(72)
            box.paragraph_format.space_after = Pt(72)
            run = box.add_run(f"[FIGURE S{number} LAYOUT PLACEHOLDER]\nNOT FOR SUBMISSION")
            main._set_font(run, size=14, bold=True, color="A61B1B")
            continue
        with Image.open(image_path) as image:
            width_px, height_px = image.size
        width_inches = 6.5
        height_inches = width_inches * height_px / width_px
        if height_inches > 8.2:
            height_inches = 8.2
            width_inches = height_inches * width_px / height_px
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
        main._set_picture_alt(shape, f"Figure S{number}: {title}. See the supplementary legend for panel-level description.")


def build_supplementary_docx(
    *,
    mode: str,
    supplement: str | Path,
    output: str | Path,
    main_manuscript: str | Path | None = None,
    main_compile_receipt: str | Path | None = None,
    figure_summary: str | Path | None = None,
    receipt: str | Path | None = None,
) -> dict[str, Any]:
    _require(mode in {"final", "layout-fixture"}, "unsupported supplementary DOCX mode")
    final = mode == "final"
    source_path = Path(supplement).resolve()
    output_path = Path(output).resolve()
    receipt_path = Path(receipt).resolve() if receipt else output_path.with_suffix(output_path.suffix + ".receipt.json")
    _require(output_path.suffix.casefold() == ".docx", "output must use .docx")
    _require(output_path != receipt_path, "DOCX and receipt paths must differ")
    _require(not output_path.exists() and not receipt_path.exists(), "refusing to overwrite output or receipt")
    supplement_raw, supplement_text = _read_text(source_path, "supplement")
    frontmatter = _extract_frontmatter(supplement_text)

    main_raw = b""
    compile_raw = b""
    compile_payload: dict[str, Any] = {}
    figure_raw = b""
    figure_payload: dict[str, Any] = {}
    figures: list[tuple[int, str, Path]] = []
    table_bundle: dict[str, Any] | None = None
    if final:
        _require(all(value is not None for value in (main_manuscript, main_compile_receipt, figure_summary)), "final mode requires main manuscript/compiler receipt and supplementary figure suite")
        main_raw, main_text = _read_text(Path(main_manuscript).resolve(), "compiled main manuscript")
        compile_raw, compile_payload = _read_json(Path(main_compile_receipt).resolve(), "main compile receipt")
        _validate_final_inputs(
            supplement_text=supplement_text,
            main_text=main_text,
            main_raw=main_raw,
            compile_raw=compile_raw,
            compile_payload=compile_payload,
        )
        figure_path = Path(figure_summary).resolve()
        figure_raw, figure_payload = _read_json(figure_path, "supplementary figure suite")
        figures, table_bundle = _validate_figure_suite(
            figure_payload,
            summary_path=figure_path,
            compile_payload=compile_payload,
        )
    else:
        _require(all(value is None for value in (main_manuscript, main_compile_receipt, figure_summary)), "layout fixture must not consume final evidence")

    document = Document()
    main._configure_styles(document)
    document.core_properties.title = frontmatter["title"]
    document.core_properties.subject = "Plant Phenomics anonymized supplementary materials"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.created = datetime(2000, 1, 1, tzinfo=timezone.utc)
    document.core_properties.modified = datetime(2000, 1, 1, tzinfo=timezone.utc)
    first = document.sections[0]
    main._set_page_geometry(first, landscape=False, line_numbers=False)
    first.different_first_page_header_footer = True
    main._configure_header_footer(
        first,
        running_title="PHAxis supplementary materials",
        fixture=not final,
        hide_on_first_page=True,
    )
    _add_title_page(document, frontmatter=frontmatter, fixture=not final)
    _, number_id = main._numbering(document)
    table_count = _add_body(
        document,
        text=supplement_text,
        running_title="PHAxis supplementary materials",
        fixture=not final,
        number_id=number_id,
    )
    if table_bundle is not None:
        _add_table_data_index(document, table_bundle)
        table_count += 1
    _add_figure_plates(
        document,
        figures=figures,
        running_title="PHAxis supplementary materials",
        fixture=not final,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    staging_dir = Path(tempfile.mkdtemp(prefix=".supp-docx-", dir=output_path.parent))
    # Keep private staging components short.  The final journal-facing name is
    # applied only by the no-overwrite publication step, which avoids repeating
    # a long DOCX/receipt basename below the staging directory on Windows.
    temporary_docx = staging_dir / "document.docx"
    try:
        document.save(temporary_docx)
        main._normalize_docx_zip(temporary_docx)
        result: dict[str, Any] = {
            "schema_version": SCHEMA_VERSION,
            "status": "completed_final_anonymized_supplementary_docx" if final else "completed_anonymized_supplementary_layout_fixture_not_for_submission",
            "mode": mode,
            "submission_use_allowed": final,
            "target_journal": "Plant Phenomics",
            "software_version": frontmatter["software_version"],
            "supplement_sha256": hashlib.sha256(supplement_raw).hexdigest(),
            "main_manuscript_sha256": hashlib.sha256(main_raw).hexdigest() if final else None,
            "main_compile_receipt_sha256": hashlib.sha256(compile_raw).hexdigest() if final else None,
            "main_compile_receipt_identity_sha256": compile_payload.get("receipt_identity_sha256") if final else None,
            "submission_metadata_consumed": False,
            "figure_summary_sha256": hashlib.sha256(figure_raw).hexdigest() if final else None,
            "supplementary_figure_suite_identity_sha256": figure_payload.get("supplementary_figure_bundle_identity_sha256") if final else None,
            "docx_sha256": main._sha256_file(temporary_docx),
            "reviewer_visible": True,
            "anonymized_supplement_separate": True,
            "anonymous_core_creator_empty": True,
            "supplementary_figure_count": 9,
            "supplementary_table_data_resource_count": 10,
            "supplementary_table_data_materialized": final,
            "supplementary_table_bundle_receipt_sha256": (
                table_bundle["receipt_sha256"] if table_bundle is not None else None
            ),
            "supplementary_table_bundle_identity_sha256": (
                table_bundle["bundle_identity_sha256"]
                if table_bundle is not None
                else None
            ),
            "supplementary_table_item_identity_sha256": (
                {
                    stem: record["item_identity_sha256"]
                    for stem, record in table_bundle["items"].items()
                }
                if table_bundle is not None
                else None
            ),
            "embedded_markdown_table_count": table_count,
            "style_contract": main.STYLE_CONTRACT,
            "blind_images_used": 0,
            "canonical_annotations_read": False,
            "root_cap_region_statistics_included": False,
        }
        if final:
            result.update(
                {
                    "model_bundle_id": compile_payload["model_bundle_id"],
                    "root_expert_id": compile_payload["root_expert_id"],
                    "hair_identity_count_expert": compile_payload["hair_identity_count_expert"],
                }
            )
        result["receipt_identity_sha256"] = main._canonical_hash(result)
        temporary_receipt = staging_dir / "receipt.json"
        temporary_receipt.write_text(
            json.dumps(result, ensure_ascii=False, allow_nan=False, indent=2, sort_keys=True) + "\n",
            encoding="utf-8",
            newline="\n",
        )
        main._publish_transaction(
            ((temporary_docx, output_path), (temporary_receipt, receipt_path))
        )
        return result
    finally:
        for child in staging_dir.iterdir() if staging_dir.exists() else []:
            child.unlink(missing_ok=True)
        staging_dir.rmdir()


def _parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--mode", choices=("final", "layout-fixture"), required=True)
    parser.add_argument("--supplement", type=Path, required=True)
    parser.add_argument("--main-manuscript", type=Path)
    parser.add_argument("--main-compile-receipt", type=Path)
    parser.add_argument("--figure-summary", type=Path)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--receipt", type=Path)
    return parser


def main_cli(argv: Sequence[str] | None = None) -> int:
    args = _parser().parse_args(argv)
    result = build_supplementary_docx(
        mode=args.mode,
        supplement=args.supplement,
        output=args.output,
        main_manuscript=args.main_manuscript,
        main_compile_receipt=args.main_compile_receipt,
        figure_summary=args.figure_summary,
        receipt=args.receipt,
    )
    print(result["receipt_identity_sha256"])
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main_cli())
    except SupplementaryDocxError as error:
        print(f"blocked: {error}", file=sys.stderr)
        raise SystemExit(2)
