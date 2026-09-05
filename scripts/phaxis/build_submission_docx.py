#!/usr/bin/env python3
"""Build the PHAxis Plant Phenomics double-anonymous submission bundle.

``final`` mode accepts only a placeholder-free compiled manuscript, its
hash-closed compiler receipt, a final six-figure suite, and author-verified
title-page metadata.  It emits a separate editor-only title page and a
reviewer-visible anonymous main manuscript in one fail-closed transaction.
``layout-fixture`` mode exists only for typography and page-layout testing; it
is visibly marked and can never authorize submission.

The builder does not discover evidence, read images or annotations, or start a
GPU program.  It uses the ``narrative_proposal`` document-skill preset with one
named ``plant_phenomics_manuscript`` override: restrained Times New Roman
typography, continuous line numbers, and journal-style tables without vertical
rules.  All geometry is encoded explicitly in OOXML.
"""

from __future__ import annotations

import argparse
from copy import deepcopy
from datetime import datetime, timezone
import hashlib
import io
import json
import os
from pathlib import Path
import re
import sys
import tempfile
from typing import Any, Iterable, Mapping, Sequence
import uuid
import zipfile

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


PROJECT_ROOT = Path(__file__).resolve().parents[2]
if str(PROJECT_ROOT / "src") not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT / "src"))

from phaxis.manuscript_contract import (  # noqa: E402
    ABSTRACT_WORD_LIMIT,
    ManuscriptTextContractError,
    require_abstract_within_limit,
)


SCHEMA_VERSION = "PHAxis-submission-docx-build-2.0"
METADATA_SCHEMA = "PHAxis-submission-title-metadata-2.0"
COMPILE_RECEIPT_SCHEMA = "PHAxis-manuscript-compile-receipt-1.2"
FIGURE_SUITE_SCHEMA = "PHAxis-publication-figure-suite-1.0"
FINAL_COMPILE_STATUS = "completed_strict_final_manuscript_compilation"
FINAL_FIGURE_STATUS = "final_sealed_strict_train399_only"
FINAL_METADATA_STATUS = "complete_author_verified_submission_metadata"
LAYOUT_METADATA_STATUS = "layout_fixture_not_author_verified"
PLACEHOLDER = re.compile(r"\{\{[A-Z0-9_]+\}\}")
PLACEMENT_MARKER = re.compile(r"^`?\[PLACE (?:FIGURE|TABLE)\b.*\]`?$", re.I)
HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
TABLE_SEPARATOR_CELL = re.compile(r"^:?-{3,}:?$")
REFERENCE_LINE = re.compile(r"^(\d+)\.\s+(.+)$")
INTERNAL_SECTION = "## 16. Machine-Fill Placeholder Registry"
DOUBLE_ANONYMOUS_ARTICLE_TYPE = "Software and Hardware Article"
EDITOR_ONLY_DECLARATIONS = (
    "Acknowledgments",
    "Funding",
    "Author Contributions",
    "Competing Interests",
)
FIGURE_STEMS = (
    "Figure_01_biological_measurement_design",
    "Figure_02_train399_development_evidence",
    "Figure_03_measurement_assurance",
    "Figure_04_difficult_image_interpretability",
    "Figure_05_exploratory_phenotype_atlas",
    "Figure_06_reproducibility_and_efficiency",
)


STYLE_CONTRACT: dict[str, Any] = {
    "base_preset": "narrative_proposal",
    "named_override": "plant_phenomics_manuscript",
    "page": {
        "portrait_inches": [8.5, 11.0],
        "landscape_inches": [11.0, 8.5],
        "margins_inches": {"top": 1.0, "right": 1.0, "bottom": 1.0, "left": 1.0},
        "header_footer_distance_inches": 0.492,
        "portrait_content_width_dxa": 9360,
        "landscape_content_width_dxa": 12960,
    },
    "body": {
        "font": "Times New Roman",
        "size_pt": 11.0,
        "alignment": "left",
        "before_pt": 0.0,
        "after_pt": 6.0,
        "line_spacing": 1.5,
    },
    "headings": {
        "h1": {"size_pt": 14.0, "color": "000000", "before_pt": 12.0, "after_pt": 6.0},
        "h2": {"size_pt": 12.0, "color": "000000", "before_pt": 10.0, "after_pt": 5.0},
        "h3": {"size_pt": 11.0, "color": "000000", "before_pt": 8.0, "after_pt": 4.0},
    },
    "lists": {
        "marker_aligned_at_inches": 0.181,
        "text_indent_at_inches": 0.375,
        "hanging_inches": 0.194,
        "after_pt": 4.0,
        "line_spacing": 1.208,
    },
    "tables": {
        "portrait_width_dxa": 9360,
        "landscape_width_dxa": 12960,
        "indent_dxa": 120,
        "cell_margins_dxa": {"top": 80, "bottom": 80, "start": 120, "end": 120},
        "font_size_pt": 8.0,
        "header_fill": "F4F6F9",
        "vertical_rules": False,
    },
    "line_numbering": {"count_by": 1, "distance_dxa": 360, "restart": "continuous"},
}


class SubmissionDocxError(RuntimeError):
    """A source, layout, or publication gate is invalid."""


def _require(condition: bool, message: str) -> None:
    if not condition:
        raise SubmissionDocxError(message)


def _is_sha256(value: Any) -> bool:
    return (
        isinstance(value, str)
        and len(value) == 64
        and all(character in "0123456789abcdef" for character in value)
    )


def _canonical_bytes(payload: Any) -> bytes:
    try:
        return json.dumps(
            payload,
            ensure_ascii=False,
            allow_nan=False,
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
    except (TypeError, ValueError) as error:
        raise SubmissionDocxError("payload is not finite canonical JSON") from error


def _canonical_hash(payload: Any) -> str:
    return hashlib.sha256(_canonical_bytes(payload)).hexdigest()


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(8 * 1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _figure_suite_identity_preimage(
    *,
    status: str,
    figure_hashes: Mapping[str, Any],
    source_hashes: Mapping[str, str],
    figure_input_assembly_identity_sha256: str,
    model_contract_proposal_identity_sha256: str,
    model_contract_public_identity: Mapping[str, str],
    train399_prediction_input_provenance: Mapping[str, Any],
    supplementary_table_bundle_identity_sha256: str,
    supplementary_table_bundle_receipt_sha256: str,
) -> dict[str, Any]:
    """Mirror the small dependency-free canonical publication preimage."""
    return {
        "status": str(status),
        "figure_hashes": deepcopy(dict(figure_hashes)),
        "source_hashes": dict(source_hashes),
        "figure_input_assembly_identity_sha256": str(
            figure_input_assembly_identity_sha256
        ),
        "model_contract_proposal_identity_sha256": str(
            model_contract_proposal_identity_sha256
        ),
        "model_contract_public_identity": dict(model_contract_public_identity),
        "train399_prediction_input_provenance": deepcopy(
            dict(train399_prediction_input_provenance)
        ),
        "supplementary_table_bundle_identity_sha256": str(
            supplementary_table_bundle_identity_sha256
        ),
        "supplementary_table_bundle_receipt_sha256": str(
            supplementary_table_bundle_receipt_sha256
        ),
    }


def _strict_pairs(pairs: list[tuple[str, Any]]) -> dict[str, Any]:
    result: dict[str, Any] = {}
    for key, value in pairs:
        if key in result:
            raise ValueError(f"duplicate JSON key: {key}")
        result[key] = value
    return result


def _reject_constant(value: str) -> None:
    raise ValueError(f"non-finite JSON constant: {value}")


def _read_json(path: Path, role: str) -> tuple[bytes, dict[str, Any]]:
    _require(not path.is_symlink(), f"{role} may not be a symlink")
    _require(path.is_file(), f"{role} does not exist: {path}")
    raw = path.read_bytes()
    try:
        payload = json.loads(
            raw.decode("utf-8"),
            object_pairs_hook=_strict_pairs,
            parse_constant=_reject_constant,
        )
    except (UnicodeError, json.JSONDecodeError, ValueError) as error:
        raise SubmissionDocxError(f"{role} is not strict UTF-8 JSON") from error
    _require(isinstance(payload, dict), f"{role} must contain one JSON object")
    return raw, payload


def _verify_seal(payload: Mapping[str, Any], field: str, role: str) -> None:
    identity = payload.get(field)
    _require(_is_sha256(identity), f"{role} has no valid {field}")
    unsigned = deepcopy(dict(payload))
    unsigned.pop(field, None)
    _require(_canonical_hash(unsigned) == identity, f"{role} identity seal mismatch")


def _read_text(path: Path, role: str) -> tuple[bytes, str]:
    _require(not path.is_symlink(), f"{role} may not be a symlink")
    _require(path.is_file(), f"{role} does not exist: {path}")
    raw = path.read_bytes()
    try:
        text = raw.decode("utf-8")
    except UnicodeError as error:
        raise SubmissionDocxError(f"{role} must be UTF-8") from error
    return raw, text


def _extract_frontmatter(text: str) -> dict[str, str]:
    lines = text.splitlines()
    _require(lines and lines[0].startswith("# "), "manuscript has no H1 title")
    result = {"title": lines[0][2:].strip()}
    labels = {
        "Running title": "running_title",
        "Article type": "article_type",
        "Software version": "software_version",
    }
    for line in lines[1:12]:
        match = re.match(r"^\*\*(.+?):\*\*\s*(.*?)\s*$", line)
        if match and match.group(1) in labels:
            result[labels[match.group(1)]] = match.group(2)
    for field in ("title", "running_title", "article_type", "software_version"):
        _require(bool(result.get(field)), f"manuscript frontmatter lacks {field}")
    return result


def _validate_metadata(
    payload: Mapping[str, Any], *, frontmatter: Mapping[str, str], final: bool
) -> str:
    _require(payload.get("schema_version") == METADATA_SCHEMA, "submission metadata schema changed")
    expected_status = FINAL_METADATA_STATUS if final else LAYOUT_METADATA_STATUS
    _require(payload.get("status") == expected_status, "submission metadata status is invalid")
    _verify_seal(payload, "metadata_identity_sha256", "submission metadata")
    _require(payload.get("journal") == "Plant Phenomics", "target journal changed")
    if final:
        _require(
            frontmatter["article_type"] == DOUBLE_ANONYMOUS_ARTICLE_TYPE,
            "Plant Phenomics article type is not Software and Hardware Article",
        )
    for field in ("manuscript_title", "running_title", "article_type", "software_version"):
        _require(
            payload.get(field) == frontmatter[field.replace("manuscript_", "")],
            f"submission metadata/manuscript {field} mismatch",
        )
    affiliations = payload.get("affiliations")
    authors = payload.get("authors")
    _require(isinstance(affiliations, list) and affiliations, "affiliations are missing")
    _require(isinstance(authors, list) and authors, "authors are missing")
    affiliation_ids: set[int] = set()
    for row in affiliations:
        _require(isinstance(row, Mapping), "affiliation row is invalid")
        identifier = row.get("id")
        _require(isinstance(identifier, int) and identifier > 0, "affiliation id is invalid")
        _require(identifier not in affiliation_ids, "duplicate affiliation id")
        _require(isinstance(row.get("text"), str) and row["text"].strip(), "affiliation text is empty")
        affiliation_ids.add(identifier)
    corresponding = 0
    for row in authors:
        _require(isinstance(row, Mapping), "author row is invalid")
        _require(isinstance(row.get("full_name"), str) and row["full_name"].strip(), "author name is empty")
        ids = row.get("affiliation_ids")
        _require(isinstance(ids, list) and ids, "author affiliations are missing")
        _require(all(isinstance(value, int) and value in affiliation_ids for value in ids), "author affiliation is unknown")
        _require(len(ids) == len(set(ids)), "author affiliation is duplicated")
        is_corresponding = row.get("corresponding_author")
        _require(isinstance(is_corresponding, bool), "corresponding-author flag is invalid")
        if is_corresponding:
            corresponding += 1
            _require(isinstance(row.get("email"), str) and "@" in row["email"], "corresponding author email is invalid")
            _require(
                isinstance(row.get("postal_address"), str)
                and bool(str(row["postal_address"]).strip()),
                "corresponding author postal address is empty",
            )
            _require(
                row.get("postal_address_author_verified") is True,
                "corresponding author postal address is not author-verified",
            )
        else:
            if row.get("email") not in (None, ""):
                _require(isinstance(row.get("email"), str) and "@" in row["email"], "author email is invalid")
            _require(
                row.get("postal_address") in (None, ""),
                "non-corresponding author has an unscoped postal address",
            )
            _require(
                row.get("postal_address_author_verified") in (None, False),
                "non-corresponding author postal verification flag is invalid",
            )
        orcid = row.get("orcid")
        if orcid not in (None, ""):
            _require(bool(re.fullmatch(r"\d{4}-\d{4}-\d{4}-\d{3}[\dX]", str(orcid))), "author ORCID is invalid")
    _require(corresponding >= 1, "at least one corresponding author is required")
    if final:
        serialized = json.dumps(payload, ensure_ascii=False).casefold()
        for marker in ("todo", "tbd", "placeholder", "author name", "affiliation text"):
            _require(marker not in serialized, f"final submission metadata contains marker: {marker}")
    return str(payload["metadata_identity_sha256"])


def submission_metadata_template(*, frontmatter: Mapping[str, str]) -> dict[str, Any]:
    """Return an intentionally incomplete, non-authorized metadata template."""
    payload: dict[str, Any] = {
        "schema_version": METADATA_SCHEMA,
        "status": "incomplete_template_not_for_submission",
        "journal": "Plant Phenomics",
        "manuscript_title": frontmatter["title"],
        "running_title": frontmatter["running_title"],
        "article_type": frontmatter["article_type"],
        "software_version": frontmatter["software_version"],
        "authors": [
            {
                "full_name": None,
                "affiliation_ids": [1],
                "corresponding_author": True,
                "email": None,
                "orcid": None,
                "postal_address": None,
                "postal_address_author_verified": False,
            }
        ],
        "affiliations": [{"id": 1, "text": None}],
    }
    payload["metadata_identity_sha256"] = _canonical_hash(payload)
    return payload


def _validate_compile_receipt(
    payload: Mapping[str, Any],
    *,
    receipt_raw: bytes,
    manuscript_raw: bytes,
    manuscript_text: str,
) -> None:
    _require(payload.get("schema_version") == COMPILE_RECEIPT_SCHEMA, "compile receipt schema changed")
    _require(payload.get("status") == FINAL_COMPILE_STATUS, "manuscript is not final-compiled")
    _verify_seal(payload, "receipt_identity_sha256", "compile receipt")
    _require(payload.get("output_sha256") == hashlib.sha256(manuscript_raw).hexdigest(), "compile receipt/manuscript SHA mismatch")
    _require(payload.get("unresolved_token_count") == 0, "compile receipt retains unresolved tokens")
    _require(payload.get("author_metadata_complete") is True, "compiled author statements are incomplete")
    try:
        observed_abstract_words = require_abstract_within_limit(manuscript_text)
    except ManuscriptTextContractError as error:
        raise SubmissionDocxError(str(error)) from error
    _require(
        payload.get("abstract_word_count") == observed_abstract_words,
        "compile receipt/manuscript abstract word-count mismatch",
    )
    _require(
        payload.get("abstract_word_limit") == ABSTRACT_WORD_LIMIT
        and payload.get("abstract_word_limit_passed") is True,
        "compile receipt lacks the Plant Phenomics abstract-limit Gate",
    )
    _require(payload.get("blind_images_used") == 0, "compile receipt is blind-tainted")
    _require(payload.get("root_cap_region_statistics_included") is False, "compile receipt includes root-cap-region statistics")
    _require(_is_sha256(hashlib.sha256(receipt_raw).hexdigest()), "compile receipt cannot be hashed")


def _child_file(path: Path, parent: Path, role: str) -> Path:
    _require(not path.is_symlink(), f"{role} may not be a symlink")
    resolved = path.resolve()
    try:
        resolved.relative_to(parent.resolve())
    except ValueError as error:
        raise SubmissionDocxError(f"{role} is outside the figure suite") from error
    _require(resolved.is_file(), f"{role} is missing")
    return resolved


def _validate_figure_summary(
    payload: Mapping[str, Any], *, summary_path: Path, compile_receipt: Mapping[str, Any]
) -> list[tuple[int, str, Path]]:
    _require(payload.get("schema_version") == FIGURE_SUITE_SCHEMA, "figure-suite schema changed")
    _require(payload.get("status") == FINAL_FIGURE_STATUS, "figure suite is not final")
    _require(payload.get("submission_use_allowed") is True, "figure suite forbids submission")
    _require(payload.get("formal_train399_only_gate_passed") is True, "figure suite lacks train399 Gate")
    _require(payload.get("blind_images_used") == 0, "figure suite is blind-tainted")
    claim = payload.get("claim_contract")
    _require(isinstance(claim, Mapping), "figure-suite claim contract is missing")
    _require(claim.get("main_figure_count") == 6, "figure-suite count changed")
    _require(claim.get("root_cap_region_statistics_included") is False, "figure suite includes root-cap-region statistics")
    _require(claim.get("canonical_annotations_read") is False, "figure suite used canonical annotations during deployment")
    for left, right, label in (
        (payload.get("model_contract_proposal_identity_sha256"), compile_receipt.get("model_contract_proposal_identity_sha256"), "proposal identity"),
        (payload.get("figure_input_assembly_identity_sha256"), compile_receipt.get("figure_input_assembly_identity_sha256"), "figure-input assembly"),
        (payload.get("model_bundle_id"), compile_receipt.get("model_bundle_id"), "model bundle"),
        (payload.get("root_expert_id"), compile_receipt.get("root_expert_id"), "root expert"),
        (payload.get("hair_identity_expert_id"), compile_receipt.get("hair_identity_count_expert"), "hair expert"),
    ):
        _require(left == right, f"figure suite/compiler {label} mismatch")
    identity = payload.get("figure_suite_identity_sha256")
    _require(_is_sha256(identity), "figure-suite identity is invalid")
    expected = _canonical_hash(
        _figure_suite_identity_preimage(
            status="final",
            figure_hashes=payload.get("figure_bundle_sha256"),
            source_hashes=payload.get("source_summary_sha256"),
            figure_input_assembly_identity_sha256=payload.get("figure_input_assembly_identity_sha256"),
            model_contract_proposal_identity_sha256=payload.get("model_contract_proposal_identity_sha256"),
            model_contract_public_identity=payload.get("model_contract_public_identity"),
            train399_prediction_input_provenance=payload.get("train399_prediction_input_provenance"),
            supplementary_table_bundle_identity_sha256=payload.get(
                "supplementary_table_bundle_identity_sha256"
            ),
            supplementary_table_bundle_receipt_sha256=payload.get(
                "supplementary_table_bundle_receipt_sha256"
            ),
        )
    )
    _require(expected == identity, "figure-suite identity seal mismatch")
    figures = payload.get("figures")
    bundle_hashes = payload.get("figure_bundle_sha256")
    _require(isinstance(figures, Mapping) and list(figures) == list(FIGURE_STEMS), "figure order/set changed")
    _require(isinstance(bundle_hashes, Mapping), "figure bundle hashes are missing")
    result: list[tuple[int, str, Path]] = []
    parent = summary_path.resolve().parent
    for expected_number, stem in enumerate(FIGURE_STEMS, start=1):
        record = figures[stem]
        _require(isinstance(record, Mapping), f"figure {expected_number} record is invalid")
        _require(record.get("number") == expected_number and record.get("status") == "final", f"figure {expected_number} is not final")
        bundle = record.get("bundle")
        hashes = bundle_hashes.get(stem)
        _require(isinstance(bundle, Mapping) and isinstance(hashes, Mapping), f"figure {expected_number} bundle is invalid")
        files = bundle.get("files")
        _require(isinstance(files, Mapping) and isinstance(files.get("png"), str), f"figure {expected_number} PNG is missing")
        image_path = _child_file(Path(files["png"]), parent, f"figure {expected_number} PNG")
        _require(_sha256_file(image_path) == hashes.get("png"), f"figure {expected_number} PNG hash mismatch")
        result.append((expected_number, str(record.get("title", stem)), image_path))
    return result


def _set_font(run, *, name: str = "Times New Roman", size: float | None = None, bold: bool | None = None, italic: bool | None = None, color: str = "000000") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def _style_font(style, name: str, size: float, color: str = "000000") -> None:
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)


def _configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    _style_font(normal, "Times New Roman", 11.0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for name, key, bold, italic in (
        ("Heading 1", "h1", True, False),
        ("Heading 2", "h2", True, False),
        ("Heading 3", "h3", True, True),
    ):
        token = STYLE_CONTRACT["headings"][key]
        style = styles[name]
        _style_font(style, "Times New Roman", token["size_pt"], token["color"])
        style.font.bold = bold
        style.font.italic = italic
        style.paragraph_format.space_before = Pt(token["before_pt"])
        style.paragraph_format.space_after = Pt(token["after_pt"])
        style.paragraph_format.keep_with_next = True
    for name in ("Title", "Subtitle", "Caption"):
        _style_font(styles[name], "Times New Roman", 11.0)
    styles["Title"].font.size = Pt(16)
    styles["Title"].font.bold = True
    styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Title"].paragraph_format.space_after = Pt(18)
    styles["Subtitle"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Caption"].font.bold = False
    styles["Caption"].font.italic = False
    styles["Caption"].paragraph_format.space_before = Pt(4)
    styles["Caption"].paragraph_format.space_after = Pt(6)
    styles["Caption"].paragraph_format.line_spacing = 1.15


def _set_page_geometry(section, *, landscape: bool, line_numbers: bool) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.0)
        section.page_height = Inches(8.5)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
    for field in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, field, Inches(1.0))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    sect_pr = section._sectPr
    for existing in list(sect_pr.findall(qn("w:lnNumType"))):
        sect_pr.remove(existing)
    if line_numbers:
        element = OxmlElement("w:lnNumType")
        element.set(qn("w:countBy"), "1")
        element.set(qn("w:distance"), "360")
        element.set(qn("w:restart"), "continuous")
        sect_pr.append(element)


def _page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)
    _set_font(run, size=9, color="666666")


def _configure_header_footer(
    section,
    *,
    running_title: str,
    fixture: bool,
    hide_on_first_page: bool = False,
) -> None:
    # python-docx copies section properties when a section is added.  Without an
    # explicit reset, the title-page ``different first page`` flag propagates to
    # the first body, table, reference, and figure page and silently removes its
    # running header/page number in Word.  Only the manuscript title page is
    # intentionally unnumbered.
    section.different_first_page_header_footer = hide_on_first_page
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text = "LAYOUT FIXTURE - NOT FOR SUBMISSION" if fixture else running_title
    run = paragraph.add_run(text)
    _set_font(run, size=8.5, bold=fixture, color="A61B1B" if fixture else "666666")
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _page_field(paragraph)


def _numbering(document: Document) -> tuple[int, int]:
    numbering = document.part.numbering_part.element
    existing_abstract = [int(element.get(qn("w:abstractNumId"))) for element in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(element.get(qn("w:numId"))) for element in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abstract, default=0) + 1
    num_id = max(existing_num, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    # Reserve enough hanging-indent width for two-digit reference labels.
    # At the previous 540/280 twip setting Word exhausted the number tab at
    # item 10 and rendered labels such as ``10.Tsang`` without visible
    # separation.  A 0.5-inch text start remains compact while keeping 1--20
    # visually distinct from the reference body.
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "20")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend((tabs, ind, spacing))
    level.extend((start, num_fmt, level_text, suffix, p_pr))
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return abstract_id, num_id


def _apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num))


def _suppress_line_number(paragraph) -> None:
    """Prevent layout-only spacer paragraphs from receiving a line number."""

    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:suppressLineNumbers")) is None:
        p_pr.append(OxmlElement("w:suppressLineNumbers"))


def _add_inline(paragraph, text: str, *, base_size: float = 11.0) -> None:
    token = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*|\[[^\]]+\]\([^\)]+\))")
    cursor = 0
    for match in token.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            _set_font(run, size=base_size)
        item = match.group(0)
        if item.startswith("**"):
            run = paragraph.add_run(item[2:-2])
            _set_font(run, size=base_size, bold=True)
        elif item.startswith("`"):
            run = paragraph.add_run(item[1:-1])
            _set_font(run, name="Courier New", size=max(8.0, base_size - 1.0))
        elif item.startswith("*"):
            run = paragraph.add_run(item[1:-1])
            _set_font(run, size=base_size, italic=True)
        else:
            label, url = re.match(r"^\[([^\]]+)\]\(([^\)]+)\)$", item).groups()  # type: ignore[union-attr]
            run = paragraph.add_run(f"{label} ({url})")
            _set_font(run, size=base_size)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        _set_font(run, size=base_size)


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip()
    _require(stripped.startswith("|") and stripped.endswith("|"), "malformed Markdown table row")
    return [cell.strip() for cell in stripped[1:-1].split("|")]


def _table_blocks(lines: list[str]) -> list[tuple[str, Any]]:
    blocks: list[tuple[str, Any]] = []
    paragraph_lines: list[str] = []

    def flush() -> None:
        if paragraph_lines:
            blocks.append(("paragraph", " ".join(value.strip() for value in paragraph_lines)))
            paragraph_lines.clear()

    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line.strip():
            flush()
            index += 1
            continue
        heading = HEADING.match(line)
        if heading:
            flush()
            blocks.append(("heading", (len(heading.group(1)), heading.group(2))))
            index += 1
            continue
        if line.lstrip().startswith("|"):
            flush()
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                rows.append(_split_table_row(lines[index]))
                index += 1
            _require(len(rows) >= 2, "Markdown table has no separator")
            _require(all(TABLE_SEPARATOR_CELL.fullmatch(cell) for cell in rows[1]), "Markdown table separator is invalid")
            width = len(rows[0])
            _require(width >= 2 and all(len(row) == width for row in rows), "Markdown table column count changed")
            blocks.append(("table", [rows[0], *rows[2:]]))
            continue
        if PLACEMENT_MARKER.fullmatch(line.strip()):
            flush()
            blocks.append(("placement", line.strip()))
            index += 1
            continue
        reference = REFERENCE_LINE.match(line)
        if reference:
            flush()
            blocks.append(("reference", (int(reference.group(1)), reference.group(2))))
            index += 1
            continue
        paragraph_lines.append(line)
        index += 1
    flush()
    return blocks


def _column_widths(headers: Sequence[str], total: int) -> list[int]:
    count = len(headers)
    normalized = [header.casefold() for header in headers]
    if count == 5 and normalized[0] == "id":
        base = [600, 1600, 2700, 750, total - 5650]
    elif count == 7 and "measurement / estimand" in normalized[0]:
        base = [1500, 2000, 1500, 1400, 1300, 1300, total - 9000]
    elif count == 7 and normalized[0] == "endpoint":
        base = [1600, 1800, 1300, 1700, 900, 2600, total - 9900]
    else:
        weights = [max(6, min(35, len(header))) for header in headers]
        raw = [max(650, round(total * weight / sum(weights))) for weight in weights]
        delta = total - sum(raw)
        raw[-1] += delta
        base = raw
    _require(sum(base) == total and all(value > 0 for value in base), "table-width allocation is invalid")
    return base


def _cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in STYLE_CONTRACT["tables"]["cell_margins_dxa"].items():
        tag = "start" if edge == "start" else "end" if edge == "end" else edge
        element = margins.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.first_child_found_in("w:tblBorders")
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    for edge in ("left", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)


def _add_table(document: Document, rows: Sequence[Sequence[str]], *, landscape: bool) -> None:
    total = 12960 if landscape else 9360
    widths = _column_widths(rows[0], total)
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        element = tbl_pr.first_child_found_in(tag)
        if element is not None:
            tbl_pr.remove(element)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.extend((tbl_w, tbl_ind, layout))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    _set_table_borders(table)
    for row_index, (row, values) in enumerate(zip(table.rows, rows, strict=True)):
        # Keep each scientific result/ontology record intact.  Without
        # ``cantSplit`` Word may divide a long row at the page boundary, which
        # can strand a token/value fragment at the top of the next page and
        # detach it visually from its endpoint label.  All manuscript rows are
        # individually shorter than the printable page height, so moving the
        # complete row is both safe and substantially easier to audit.
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        tr_pr.append(cant_split)
        if row_index == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for column_index, (cell, value) in enumerate(zip(row.cells, values, strict=True)):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[column_index]))
            tc_w.set(qn("w:type"), "dxa")
            _cell_margins(cell)
            if row_index == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), STYLE_CONTRACT["tables"]["header_fill"])
                tc_pr.append(shading)
                bottom = OxmlElement("w:tcBorders")
                edge = OxmlElement("w:bottom")
                edge.set(qn("w:val"), "single")
                edge.set(qn("w:sz"), "6")
                edge.set(qn("w:color"), "000000")
                bottom.append(edge)
                tc_pr.append(bottom)
            paragraph = cell.paragraphs[0]
            paragraph.text = ""
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if column_index == 0 or len(value) <= 12
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            _add_inline(paragraph, value, base_size=8.0)
            if row_index == 0:
                # A repeating Word table header can otherwise be laid out as
                # the final row on a page before the first data row moves to
                # the next page.  Chaining every header-cell paragraph to the
                # following paragraph, together with ``cantSplit`` on the
                # first data row, keeps the initial header and first record as
                # one auditable visual unit while still allowing the header
                # to repeat on later pages.
                paragraph.paragraph_format.keep_with_next = True
                for run in paragraph.runs:
                    run.bold = True
    spacer = document.add_paragraph()
    _suppress_line_number(spacer)
    spacer.paragraph_format.space_after = Pt(4)


def _semantic_h2_title(title: str) -> str:
    """Return a heading title without a leading manuscript section number."""

    return re.sub(r"^\d+\.\s*", "", title).strip()


def _extract_editor_only_declarations(
    text: str, *, final: bool
) -> dict[str, str]:
    """Extract the four identity-bearing declarations from compiled Markdown."""

    wanted = {value.casefold(): value for value in EDITOR_ONLY_DECLARATIONS}
    collected: dict[str, list[str]] = {}
    current: str | None = None
    for line in text.splitlines():
        heading = HEADING.match(line)
        if heading and len(heading.group(1)) == 2:
            semantic = _semantic_h2_title(heading.group(2)).casefold()
            current = wanted.get(semantic)
            if current is not None:
                _require(
                    current not in collected,
                    f"duplicate editor-only declaration section: {current}",
                )
                collected[current] = []
            continue
        if current is not None:
            collected[current].append(line)
    result = {
        label: "\n".join(collected.get(label, [])).strip()
        for label in EDITOR_ONLY_DECLARATIONS
    }
    if final:
        for label, value in result.items():
            _require(bool(value), f"editor-only declaration is empty or absent: {label}")
            _require(
                PLACEHOLDER.search(value) is None,
                f"editor-only declaration retains a placeholder: {label}",
            )
    else:
        for label, value in list(result.items()):
            if not value:
                result[label] = "LAYOUT DECLARATION PLACEHOLDER - NOT FOR SUBMISSION"
    return result


def _add_title_page(
    document: Document,
    metadata: Mapping[str, Any],
    declarations: Mapping[str, str],
    *,
    fixture: bool,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run(
        "LAYOUT FIXTURE - NOT FOR SUBMISSION"
        if fixture
        else "Plant Phenomics | EDITOR-ONLY TITLE PAGE"
    )
    _set_font(run, size=11, bold=True, color="A61B1B" if fixture else "555555")
    # Do not use Word's built-in Title style here.  Some Word installations
    # retain a theme-defined bottom paragraph border on that style, producing a
    # spurious blue rule below the manuscript title even after its font is
    # restyled.  A plain paragraph with explicit journal formatting is stable
    # across Word and LibreOffice.
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(18)
    paragraph.add_run(str(metadata["manuscript_title"]))
    for run in paragraph.runs:
        _set_font(run, size=16, bold=True)
    author_paragraph = document.add_paragraph()
    author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_paragraph.paragraph_format.space_after = Pt(12)
    for index, author in enumerate(metadata["authors"]):
        if index:
            run = author_paragraph.add_run(", ")
            _set_font(run, size=11)
        run = author_paragraph.add_run(str(author["full_name"]))
        _set_font(run, size=11, bold=True)
        marker = ",".join(str(value) for value in author["affiliation_ids"])
        if author["corresponding_author"]:
            marker += "*"
        superscript = author_paragraph.add_run(marker)
        _set_font(superscript, size=8)
        superscript.font.superscript = True
    for affiliation in metadata["affiliations"]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(3)
        marker = paragraph.add_run(str(affiliation["id"]))
        _set_font(marker, size=8)
        marker.font.superscript = True
        run = paragraph.add_run(f" {affiliation['text']}")
        _set_font(run, size=9.5)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(14)
    contacts = []
    for author in metadata["authors"]:
        if not author["corresponding_author"]:
            continue
        value = (
            f"{author['full_name']} ({author['email']}); "
            f"postal address: {author['postal_address']}"
        )
        if author.get("orcid"):
            value += f"; ORCID: {author['orcid']}"
        contacts.append(value)
    run = paragraph.add_run("*Correspondence: " + " | ".join(contacts))
    _set_font(run, size=9.5)
    details = (
        ("Article type", metadata["article_type"]),
        ("Running title", metadata["running_title"]),
        ("Software version", metadata["software_version"]),
        ("Main figures", "6"),
        ("Main tables", "3"),
    )
    for label, value in details:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.75)
        paragraph.paragraph_format.right_indent = Inches(0.75)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(f"{label}: ")
        _set_font(run, size=10, bold=True)
        run = paragraph.add_run(str(value))
        _set_font(run, size=10)

    for label in EDITOR_ONLY_DECLARATIONS:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(label)
        _set_font(run, size=10.5, bold=True)
        for block in re.split(r"\n\s*\n", str(declarations[label]).strip()):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.25)
            paragraph.paragraph_format.space_after = Pt(4)
            _add_inline(paragraph, " ".join(block.splitlines()), base_size=9.5)


def _add_anonymous_title_block(
    document: Document, *, frontmatter: Mapping[str, str], fixture: bool
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(
        "LAYOUT FIXTURE - NOT FOR SUBMISSION"
        if fixture
        else "Plant Phenomics | ANONYMIZED MANUSCRIPT"
    )
    _set_font(run, size=10, bold=True, color="A61B1B" if fixture else "555555")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(str(frontmatter["title"]))
    _set_font(run, size=16, bold=True)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run(str(frontmatter["article_type"]))
    _set_font(run, size=10, italic=True, color="555555")


def _clean_body_lines(text: str, *, final: bool) -> list[str]:
    if INTERNAL_SECTION in text:
        text = text.split(INTERNAL_SECTION, 1)[0].rstrip() + "\n"
    lines = text.splitlines()
    _require(lines and lines[0].startswith("# "), "manuscript title is missing")
    cleaned: list[str] = []
    identity_titles = {value.casefold() for value in EDITOR_ONLY_DECLARATIONS}
    removed_numbers: list[int] = []
    skipping_identity_section = False
    for line in lines[1:]:
        if re.match(r"^\*\*(?:Running title|Article type|Software version|Draft status):\*\*", line):
            continue
        heading = HEADING.match(line)
        if heading and len(heading.group(1)) == 2:
            raw_title = heading.group(2)
            semantic = _semantic_h2_title(raw_title)
            numbered = re.match(r"^(\d+)\.\s*(.+)$", raw_title)
            if semantic.casefold() in identity_titles:
                skipping_identity_section = True
                if numbered:
                    removed_numbers.append(int(numbered.group(1)))
                continue
            skipping_identity_section = False
            if numbered and removed_numbers:
                original = int(numbered.group(1))
                replacement = original - sum(value < original for value in removed_numbers)
                line = f"{heading.group(1)} {replacement}. {numbered.group(2)}"
        if skipping_identity_section:
            continue
        cleaned.append(line)
    if final:
        anonymous_text = "\n".join(cleaned)
        _require(PLACEHOLDER.search(anonymous_text) is None, "final manuscript retains placeholders")
        for title in EDITOR_ONLY_DECLARATIONS:
            _require(
                title.casefold() not in anonymous_text.casefold(),
                f"anonymous manuscript retains editor-only section: {title}",
            )
    return cleaned


def _set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def _add_body(
    document: Document,
    text: str,
    *,
    running_title: str,
    fixture: bool,
    reference_num_id: int,
    start_new_page: bool = True,
) -> tuple[int, int]:
    body_lines = _clean_body_lines(text, final=not fixture)
    section = (
        document.add_section(WD_SECTION.NEW_PAGE)
        if start_new_page
        else document.sections[-1]
    )
    _set_page_geometry(section, landscape=False, line_numbers=True)
    _configure_header_footer(section, running_title=running_title, fixture=fixture)
    landscape = False
    table_count = 0
    reference_numbers: list[int] = []
    for kind, content in _table_blocks(body_lines):
        if kind == "heading":
            level, title = content
            semantic_title = _semantic_h2_title(title)
            if level == 2 and semantic_title.startswith("Main Table"):
                section = document.add_section(WD_SECTION.NEW_PAGE)
                _set_page_geometry(section, landscape=True, line_numbers=True)
                _configure_header_footer(section, running_title=running_title, fixture=fixture)
                landscape = True
            elif level == 2 and semantic_title == "References" and landscape:
                section = document.add_section(WD_SECTION.NEW_PAGE)
                _set_page_geometry(section, landscape=False, line_numbers=True)
                _configure_header_footer(section, running_title=running_title, fixture=fixture)
                landscape = False
            style = "Heading 1" if level == 2 else "Heading 2" if level == 3 else "Heading 3"
            paragraph = document.add_paragraph(style=style)
            _add_inline(paragraph, title, base_size=14 if style == "Heading 1" else 12 if style == "Heading 2" else 11)
            _set_keep_with_next(paragraph)
        elif kind == "table":
            table_count += 1
            _add_table(document, content, landscape=landscape)
        elif kind == "reference":
            number, value = content
            reference_numbers.append(number)
            paragraph = document.add_paragraph()
            # One point keeps the bibliography readable while avoiding a
            # near-empty spill page for the fixed 20-reference journal limit.
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.line_spacing = 1.0
            _apply_numbering(paragraph, reference_num_id)
            _add_inline(paragraph, value, base_size=9.0)
        elif kind == "placement":
            continue
        else:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.5
            _add_inline(paragraph, content, base_size=11.0)
    if reference_numbers:
        _require(reference_numbers == list(range(1, len(reference_numbers) + 1)), "reference numbering is not contiguous")
    return table_count, len(reference_numbers)


def _set_picture_alt(inline_shape, description: str) -> None:
    drawing = inline_shape._inline
    doc_pr = drawing.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description.split(":", 1)[0])


def _add_figure_plates(
    document: Document,
    figures: Sequence[tuple[int, str, Path]],
    *,
    running_title: str,
    fixture: bool,
) -> None:
    if not figures:
        figures = [(index, f"Layout placeholder for Figure {index}", Path()) for index in range(1, 7)]
    section = document.add_section(WD_SECTION.NEW_PAGE)
    _set_page_geometry(section, landscape=False, line_numbers=False)
    _configure_header_footer(section, running_title=running_title, fixture=fixture)
    heading = document.add_paragraph(style="Heading 1")
    _add_inline(heading, "Main Figures", base_size=14)
    for figure_index, (number, title, image_path) in enumerate(figures):
        if figure_index:
            document.add_page_break()
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(f"Figure {number}. {title}")
        _set_font(run, size=10, bold=True)
        if fixture:
            box = document.add_paragraph()
            box.alignment = WD_ALIGN_PARAGRAPH.CENTER
            box.paragraph_format.space_before = Pt(72)
            box.paragraph_format.space_after = Pt(72)
            run = box.add_run(f"[FIGURE {number} LAYOUT PLACEHOLDER]\nNOT FOR SUBMISSION")
            _set_font(run, size=14, bold=True, color="A61B1B")
            continue
        with Image.open(image_path) as image:
            width_px, height_px = image.size
        _require(width_px >= 1000 and height_px >= 600, f"figure {number} PNG is too small for submission assembly")
        width_inches = 6.5
        height_inches = width_inches * height_px / width_px
        if height_inches > 7.8:
            height_inches = 7.8
            width_inches = height_inches * width_px / height_px
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
        _set_picture_alt(shape, f"Figure {number}: {title}. See the manuscript figure legend for panel-level description.")


def _normalize_docx_zip(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as source:
        members = [(info.filename, source.read(info.filename), info.compress_type) for info in source.infolist()]
    # Keep the transient basename short.  A journal-facing DOCX can already
    # have a long descriptive name; repeating it here crosses the legacy
    # Windows MAX_PATH boundary inside the atomic staging directory.
    temporary = path.parent / f".zip-{uuid.uuid4().hex}.tmp"
    with zipfile.ZipFile(temporary, "w") as target:
        for name, content, compression in sorted(members):
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = compression
            info.external_attr = 0o600 << 16
            target.writestr(info, content)
    os.replace(temporary, path)


def _publish_no_overwrite(source: Path, destination: Path) -> None:
    try:
        os.link(source, destination)
    except FileExistsError as error:
        raise SubmissionDocxError(f"refusing to overwrite output: {destination}") from error
    except OSError as error:
        raise SubmissionDocxError(f"atomic no-overwrite publication failed: {destination}") from error


def _publish_transaction(members: Sequence[tuple[Path, Path]]) -> None:
    """Publish several staged files create-only, rolling back every partial member."""

    _require(bool(members), "publication transaction has no members")
    destinations = [destination for _, destination in members]
    _require(
        len(destinations) == len(set(destinations)),
        "publication transaction has duplicate destinations",
    )
    _require(
        len({destination.parent.resolve() for destination in destinations}) == 1,
        "publication transaction members must share one destination directory",
    )
    for source, destination in members:
        _require(source.is_file() and not source.is_symlink(), f"staged member is invalid: {source}")
        _require(not destination.exists(), f"refusing to overwrite output: {destination}")
    published: list[Path] = []
    try:
        for source, destination in members:
            _publish_no_overwrite(source, destination)
            published.append(destination)
    except BaseException:
        for destination in reversed(published):
            destination.unlink(missing_ok=True)
        raise


def build_submission_docx(
    *,
    mode: str,
    manuscript: str | Path,
    submission_metadata: str | Path,
    title_page_output: str | Path | None = None,
    anonymized_main_output: str | Path | None = None,
    output: str | Path | None = None,
    compile_receipt: str | Path | None = None,
    figure_summary: str | Path | None = None,
    receipt: str | Path | None = None,
) -> dict[str, Any]:
    """Build and atomically publish editor-only and reviewer-visible DOCX files.

    ``output`` is retained only as a compatibility alias for
    ``anonymized_main_output``; new callers must pass both role-specific paths.
    """
    _require(mode in {"final", "layout-fixture"}, "unsupported DOCX mode")
    final = mode == "final"
    manuscript_path = Path(manuscript).resolve()
    metadata_path = Path(submission_metadata).resolve()
    _require(
        anonymized_main_output is not None or output is not None,
        "anonymized main output is required",
    )
    _require(
        not (anonymized_main_output is not None and output is not None),
        "use either anonymized_main_output or legacy output, not both",
    )
    anonymous_path = Path(
        anonymized_main_output if anonymized_main_output is not None else output
    ).resolve()
    if title_page_output is None:
        _require(
            output is not None,
            "new double-anonymous builds require title_page_output",
        )
        title_path = anonymous_path.with_name(
            f"{anonymous_path.stem}.title_page{anonymous_path.suffix}"
        )
    else:
        title_path = Path(title_page_output).resolve()
    receipt_path = (
        Path(receipt).resolve()
        if receipt
        else anonymous_path.with_suffix(anonymous_path.suffix + ".receipt.json")
    )
    outputs = (title_path, anonymous_path, receipt_path)
    _require(
        title_path.suffix.casefold() == anonymous_path.suffix.casefold() == ".docx",
        "title page and anonymous main outputs must use .docx",
    )
    _require(len(set(outputs)) == 3, "title page, anonymous main and receipt paths must differ")
    _require(
        len({path.parent.resolve() for path in outputs}) == 1,
        "double-anonymous transaction outputs must share one directory",
    )
    _require(not any(path.exists() for path in outputs), "refusing to overwrite bundle output or receipt")
    manuscript_raw, manuscript_text = _read_text(manuscript_path, "manuscript")
    frontmatter = _extract_frontmatter(manuscript_text)
    metadata_raw, metadata_payload = _read_json(metadata_path, "submission metadata")
    metadata_identity = _validate_metadata(metadata_payload, frontmatter=frontmatter, final=final)
    compile_raw = b""
    compile_payload: dict[str, Any] = {}
    figure_raw = b""
    figure_payload: dict[str, Any] = {}
    figures: list[tuple[int, str, Path]] = []
    if final:
        _require(compile_receipt is not None and figure_summary is not None, "final mode requires compiler and figure receipts")
        compile_raw, compile_payload = _read_json(Path(compile_receipt).resolve(), "compile receipt")
        _validate_compile_receipt(
            compile_payload,
            receipt_raw=compile_raw,
            manuscript_raw=manuscript_raw,
            manuscript_text=manuscript_text,
        )
        figure_path = Path(figure_summary).resolve()
        figure_raw, figure_payload = _read_json(figure_path, "figure-suite summary")
        figures = _validate_figure_summary(figure_payload, summary_path=figure_path, compile_receipt=compile_payload)
        _require(PLACEHOLDER.search(manuscript_text) is None, "final manuscript retains machine tokens")
        _require("PROVISIONAL" not in manuscript_text.upper(), "final manuscript contains provisional marker")
    else:
        _require(compile_receipt is None and figure_summary is None, "layout fixture must not consume final receipts")

    declarations = _extract_editor_only_declarations(manuscript_text, final=final)

    title_document = Document()
    _configure_styles(title_document)
    title_document.core_properties.title = f"{frontmatter['title']} - title page"
    title_document.core_properties.subject = (
        f"Plant Phenomics {DOUBLE_ANONYMOUS_ARTICLE_TYPE} editor-only title page"
    )
    title_document.core_properties.author = "; ".join(
        str(author["full_name"]) for author in metadata_payload["authors"]
    )
    title_document.core_properties.last_modified_by = "PHAxis deterministic submission builder"
    title_document.core_properties.created = datetime(2000, 1, 1, tzinfo=timezone.utc)
    title_document.core_properties.modified = datetime(2000, 1, 1, tzinfo=timezone.utc)
    title_first = title_document.sections[0]
    _set_page_geometry(title_first, landscape=False, line_numbers=False)
    title_first.different_first_page_header_footer = True
    _add_title_page(
        title_document,
        metadata_payload,
        declarations,
        fixture=not final,
    )

    anonymous_document = Document()
    _configure_styles(anonymous_document)
    anonymous_document.core_properties.title = frontmatter["title"]
    anonymous_document.core_properties.subject = (
        f"Plant Phenomics {DOUBLE_ANONYMOUS_ARTICLE_TYPE} anonymized manuscript"
    )
    anonymous_document.core_properties.author = ""
    anonymous_document.core_properties.last_modified_by = ""
    anonymous_document.core_properties.created = datetime(2000, 1, 1, tzinfo=timezone.utc)
    anonymous_document.core_properties.modified = datetime(2000, 1, 1, tzinfo=timezone.utc)
    first = anonymous_document.sections[0]
    _set_page_geometry(first, landscape=False, line_numbers=True)
    _configure_header_footer(
        first,
        running_title=frontmatter["running_title"],
        fixture=not final,
    )
    _add_anonymous_title_block(
        anonymous_document,
        frontmatter=frontmatter,
        fixture=not final,
    )
    _, reference_num_id = _numbering(anonymous_document)
    table_count, reference_count = _add_body(
        anonymous_document,
        manuscript_text,
        running_title=frontmatter["running_title"],
        fixture=not final,
        reference_num_id=reference_num_id,
        start_new_page=False,
    )
    _require(table_count == 3, f"submission manuscript must contain exactly three tables, found {table_count}")
    _require(reference_count <= 20, f"Plant Phenomics main-reference limit exceeded: {reference_count}")
    _add_figure_plates(
        anonymous_document,
        figures,
        running_title=frontmatter["running_title"],
        fixture=not final,
    )

    anonymous_path.parent.mkdir(parents=True, exist_ok=True)
    # Preserve the journal-facing destination name only at publication time.
    # Repeating it below the private staging directory can cross the legacy
    # Windows path limit, and the default receipt name is longer still.
    staging_dir = Path(tempfile.mkdtemp(prefix=".docx-", dir=anonymous_path.parent))
    temporary_title = staging_dir / "title.docx"
    temporary_anonymous = staging_dir / "anonymous.docx"
    try:
        title_document.save(temporary_title)
        anonymous_document.save(temporary_anonymous)
        _normalize_docx_zip(temporary_title)
        _normalize_docx_zip(temporary_anonymous)
        title_sha = _sha256_file(temporary_title)
        anonymous_sha = _sha256_file(temporary_anonymous)
        result: dict[str, Any] = {
            "schema_version": SCHEMA_VERSION,
            "status": (
                "completed_final_double_anonymous_submission_bundle"
                if final
                else "completed_double_anonymous_layout_fixture_not_for_submission"
            ),
            "mode": mode,
            "submission_use_allowed": final,
            "target_journal": "Plant Phenomics",
            "article_type": frontmatter["article_type"],
            "software_version": frontmatter["software_version"],
            "manuscript_sha256": hashlib.sha256(manuscript_raw).hexdigest(),
            "compile_receipt_sha256": hashlib.sha256(compile_raw).hexdigest() if final else None,
            "compile_receipt_identity_sha256": compile_payload.get("receipt_identity_sha256") if final else None,
            "figure_summary_sha256": hashlib.sha256(figure_raw).hexdigest() if final else None,
            "figure_suite_identity_sha256": figure_payload.get("figure_suite_identity_sha256") if final else None,
            "submission_metadata_sha256": hashlib.sha256(metadata_raw).hexdigest(),
            "submission_metadata_identity_sha256": metadata_identity,
            "title_page_docx_sha256": title_sha,
            "anonymized_main_docx_sha256": anonymous_sha,
            "docx_sha256": anonymous_sha,
            "title_page_separate": True,
            "anonymized_main_separate": True,
            "editor_only_declarations": list(EDITOR_ONLY_DECLARATIONS),
            "editor_only_declaration_sha256": {
                label: hashlib.sha256(declarations[label].encode("utf-8")).hexdigest()
                for label in EDITOR_ONLY_DECLARATIONS
            },
            "reviewer_visible_identity_declarations_removed": True,
            "anonymous_core_creator_empty": True,
            "main_figure_count": 6,
            "main_table_count": table_count,
            "main_reference_count": reference_count,
            "abstract_word_count": (
                compile_payload.get("abstract_word_count") if final else None
            ),
            "abstract_word_limit": ABSTRACT_WORD_LIMIT,
            "excluded_internal_section": "Machine-Fill Placeholder Registry",
            "excluded_placement_markers": True,
            "style_contract": STYLE_CONTRACT,
            "blind_images_used": 0,
            "canonical_annotations_read": False,
            "root_cap_region_statistics_included": False,
        }
        if final:
            result.update(
                {
                    "model_bundle_id": compile_payload["model_bundle_id"],
                    "root_expert_id": compile_payload["root_expert_id"],
                    "hair_identity_count_expert": compile_payload["hair_identity_count_expert"],
                }
            )
        result["receipt_identity_sha256"] = _canonical_hash(result)
        temporary_receipt = staging_dir / "receipt.json"
        temporary_receipt.write_text(
            json.dumps(result, ensure_ascii=False, allow_nan=False, indent=2, sort_keys=True) + "\n",
            encoding="utf-8",
            newline="\n",
        )
        _publish_transaction(
            (
                (temporary_title, title_path),
                (temporary_anonymous, anonymous_path),
                (temporary_receipt, receipt_path),
            )
        )
        return result
    finally:
        for child in staging_dir.iterdir() if staging_dir.exists() else []:
            child.unlink(missing_ok=True)
        staging_dir.rmdir()


def _parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--mode", choices=("final", "layout-fixture"))
    parser.add_argument("--manuscript", type=Path, required=True)
    parser.add_argument("--submission-metadata", type=Path)
    parser.add_argument("--compile-receipt", type=Path)
    parser.add_argument("--figure-summary", type=Path)
    parser.add_argument("--title-page-output", type=Path)
    parser.add_argument("--anonymized-main-output", type=Path)
    parser.add_argument(
        "--output",
        type=Path,
        help="legacy alias for --anonymized-main-output",
    )
    parser.add_argument("--receipt", type=Path)
    parser.add_argument("--write-metadata-template", type=Path)
    return parser


def main(argv: Sequence[str] | None = None) -> int:
    args = _parser().parse_args(argv)
    if args.write_metadata_template is not None:
        _, manuscript_text = _read_text(args.manuscript.resolve(), "manuscript")
        template = submission_metadata_template(frontmatter=_extract_frontmatter(manuscript_text))
        destination = args.write_metadata_template.resolve()
        destination.parent.mkdir(parents=True, exist_ok=True)
        staging = Path(
            tempfile.mkdtemp(prefix=".metadata-v2-", dir=destination.parent)
        )
        temporary = staging / "metadata.json"
        try:
            temporary.write_text(
                json.dumps(
                    template,
                    ensure_ascii=False,
                    allow_nan=False,
                    indent=2,
                    sort_keys=True,
                )
                + "\n",
                encoding="utf-8",
                newline="\n",
            )
            _publish_no_overwrite(temporary, destination)
        finally:
            temporary.unlink(missing_ok=True)
            staging.rmdir()
        print(destination)
        return 0
    _require(args.mode is not None, "--mode is required for a DOCX build")
    _require(args.submission_metadata is not None, "--submission-metadata is required for a DOCX build")
    _require(
        args.anonymized_main_output is not None or args.output is not None,
        "--anonymized-main-output is required for a DOCX build",
    )
    _require(
        args.title_page_output is not None or args.output is not None,
        "--title-page-output is required for a DOCX build",
    )
    receipt = build_submission_docx(
        mode=args.mode,
        manuscript=args.manuscript,
        submission_metadata=args.submission_metadata,
        title_page_output=args.title_page_output,
        anonymized_main_output=args.anonymized_main_output,
        output=args.output,
        compile_receipt=args.compile_receipt,
        figure_summary=args.figure_summary,
        receipt=args.receipt,
    )
    print(receipt["receipt_identity_sha256"])
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except SubmissionDocxError as error:
        print(f"blocked: {error}", file=sys.stderr)
        raise SystemExit(2)
