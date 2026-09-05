from __future__ import annotations

import hashlib
import importlib.util
import json
from pathlib import Path
import zipfile

from docx import Document
from docx.oxml.ns import qn
from PIL import Image
import pytest


PROJECT_ROOT = Path(__file__).resolve().parents[2]
SCRIPT = PROJECT_ROOT / "scripts/phaxis/build_submission_docx.py"
SPEC = importlib.util.spec_from_file_location("phaxis_submission_docx", SCRIPT)
assert SPEC and SPEC.loader
builder = importlib.util.module_from_spec(SPEC)
SPEC.loader.exec_module(builder)


def _write_json(path: Path, payload: dict) -> Path:
    path.write_text(
        json.dumps(payload, ensure_ascii=False, allow_nan=False, indent=2, sort_keys=True)
        + "\n",
        encoding="utf-8",
    )
    return path


def _metadata(path: Path, *, final: bool) -> Path:
    payload = {
        "schema_version": builder.METADATA_SCHEMA,
        "status": builder.FINAL_METADATA_STATUS if final else builder.LAYOUT_METADATA_STATUS,
        "journal": "Plant Phenomics",
        "manuscript_title": "PHAxis submission fixture",
        "running_title": "PHAxis fixture",
        "article_type": "Software and Hardware Article",
        "software_version": "PHAxis 1.0.0",
        "authors": [
            {
                "full_name": "Layout Test Author" if not final else "Verified Test Author",
                "affiliation_ids": [1],
                "corresponding_author": True,
                "email": "verified@example.org",
                "orcid": "0000-0002-1825-0097",
                "postal_address": "1 Verified Way, Test City 10000, Testland",
                "postal_address_author_verified": True,
            }
        ],
        "affiliations": [{"id": 1, "text": "Verified Test Institute"}],
    }
    payload["metadata_identity_sha256"] = builder._canonical_hash(payload)
    return _write_json(path, payload)


def _manuscript(path: Path, *, placeholders: bool) -> Path:
    value = "{{FINAL_HAIR_F1_20UM}}" if placeholders else "0.912"
    path.write_text(
        "# PHAxis submission fixture\n\n"
        "**Running title:** PHAxis fixture  \n"
        "**Article type:** Software and Hardware Article  \n"
        "**Software version:** PHAxis 1.0.0  \n"
        "**Draft status:** internal machine source\n\n"
        "## Abstract\n\n"
        f"PHAxis measures visible biological hair identities with F1 {value}.\n\n"
        "## 1. Introduction\n\n"
        "Root hairs provide a plant-environment interface.\n\n"
        "## 5. Data Availability\n\n"
        "The dataset card, annotation schema, and reviewer example input with "
        "expected output are available at https://example.org/phaxis-data under "
        "the CC-BY-4.0 license.\n\n"
        "## 6. Code Availability\n\n"
        "Installable source, documentation, environment locks, example input, "
        "and expected output are available at https://example.org/phaxis-code "
        "under the Apache-2.0 license.\n\n"
        "## 7. Acknowledgments\n\n"
        "We thank the Verified Imaging Core for technical assistance.\n\n"
        "## 8. Funding\n\n"
        "This work was supported by Verified Grant 123.\n\n"
        "## 9. Author Contributions\n\n"
        "Verified Test Author designed and performed the study.\n\n"
        "## 10. Competing Interests\n\n"
        "The author declares no competing interests.\n\n"
        "## 13. Main Figure Legends\n\n"
        "### Figure 1. Measurement design\n\n"
        "Representative figure legend.\n\n"
        "`[PLACE FIGURE 1 NEAR INTRODUCTION END]`\n\n"
        "## 14. Main Table Legends and Shells\n\n"
        "### Table 1. Ontology\n\n"
        "| ID | Trait | Definition | Unit | Observability |\n"
        "|---|---|---|---|---|\n"
        "| R01 | Axis length | Ordered geodesic | µm | Visible axis and scale |\n\n"
        "### Table 2. Assurance\n\n"
        "| Measurement / estimand | Reference | Evidence | n | Estimate | 95% uncertainty | Boundary |\n"
        "|---|---|---|---|---:|---|---|\n"
        "| Hair identity F1 | Manual identity | Development | 44 | 0.912 | 0.89-0.93 | Not independent |\n\n"
        "### Table 3. Biology\n\n"
        "| Endpoint | Coefficient | n | Ratio | q | Sensitivity | Interpretation |\n"
        "|---|---|---|---|---:|---|---|\n"
        "| Hair count | Temperature | 10/10/10/10 | 1.20 | 0.04 | Concordant | Association |\n\n"
        "`[PLACE TABLE 3 NEAR RESULTS]`\n\n"
        "## 15. References\n\n"
        "1. Test A, et al. Reference one. 2026.\n"
        "2. Test B, et al. Reference two. 2026.\n\n"
        "## 16. Machine-Fill Placeholder Registry\n\n"
        "This internal section must not enter the DOCX.\n",
        encoding="utf-8",
    )
    return path


def _compile_receipt(path: Path, manuscript: Path) -> dict:
    digest = hashlib.sha256(b"binding").hexdigest()
    abstract_words = builder.require_abstract_within_limit(
        manuscript.read_text(encoding="utf-8")
    )
    payload = {
        "schema_version": builder.COMPILE_RECEIPT_SCHEMA,
        "status": builder.FINAL_COMPILE_STATUS,
        "output_sha256": builder._sha256_file(manuscript),
        "unresolved_token_count": 0,
        "author_metadata_complete": True,
        "abstract_word_count": abstract_words,
        "abstract_word_limit": builder.ABSTRACT_WORD_LIMIT,
        "abstract_word_limit_passed": True,
        "blind_images_used": 0,
        "root_cap_region_statistics_included": False,
        "model_contract_proposal_identity_sha256": digest,
        "figure_input_assembly_identity_sha256": hashlib.sha256(b"figure-inputs").hexdigest(),
        "model_bundle_id": "PHAXIS-V1.0.0-STRICT-TRAIN399-TEST",
        "root_expert_id": "PHAxis-root-provider-TEST",
        "hair_identity_count_expert": "PHAxis-StageB-train399-five-seed",
    }
    payload["receipt_identity_sha256"] = builder._canonical_hash(payload)
    _write_json(path, payload)
    return payload


def _figure_summary(root: Path, compile_receipt: dict) -> Path:
    root.mkdir(parents=True)
    public_identity = {
        "model_bundle_id": compile_receipt["model_bundle_id"],
        "root_expert_id": compile_receipt["root_expert_id"],
    }
    figures = {}
    figure_hashes = {}
    for number, stem in enumerate(builder.FIGURE_STEMS, start=1):
        path = root / f"{stem}.png"
        Image.new("RGB", (1200, 700), (245, 247, 250)).save(path)
        image_hash = builder._sha256_file(path)
        figures[stem] = {
            "number": number,
            "title": stem.replace("_", " "),
            "status": "final",
            "bundle": {"files": {"png": str(path.resolve())}},
        }
        figure_hashes[stem] = {
            "png": image_hash,
            "pdf": hashlib.sha256(f"pdf-{number}".encode()).hexdigest(),
            "tiff": hashlib.sha256(f"tiff-{number}".encode()).hexdigest(),
            "source_data": {},
        }
    sources = {"stageb": hashlib.sha256(b"stageb").hexdigest()}
    provenance = {"ordered_file_set_identity_sha256": hashlib.sha256(b"ordered").hexdigest()}
    supplementary_table_bundle_identity = hashlib.sha256(
        b"supplementary-table-bundle"
    ).hexdigest()
    supplementary_table_bundle_receipt = hashlib.sha256(
        b"supplementary-table-receipt"
    ).hexdigest()
    payload = {
        "schema_version": builder.FIGURE_SUITE_SCHEMA,
        "status": builder.FINAL_FIGURE_STATUS,
        "submission_use_allowed": True,
        "formal_train399_only_gate_passed": True,
        "blind_images_used": 0,
        "claim_contract": {
            "main_figure_count": 6,
            "root_cap_region_statistics_included": False,
            "canonical_annotations_read": False,
        },
        "model_contract_proposal_identity_sha256": compile_receipt[
            "model_contract_proposal_identity_sha256"
        ],
        "figure_input_assembly_identity_sha256": compile_receipt[
            "figure_input_assembly_identity_sha256"
        ],
        "model_bundle_id": compile_receipt["model_bundle_id"],
        "root_expert_id": compile_receipt["root_expert_id"],
        "hair_identity_expert_id": compile_receipt["hair_identity_count_expert"],
        "model_contract_public_identity": public_identity,
        "train399_prediction_input_provenance": provenance,
        "supplementary_table_bundle_identity_sha256": (
            supplementary_table_bundle_identity
        ),
        "supplementary_table_bundle_receipt_sha256": (
            supplementary_table_bundle_receipt
        ),
        "source_summary_sha256": sources,
        "figure_bundle_sha256": figure_hashes,
        "figures": figures,
    }
    payload["figure_suite_identity_sha256"] = builder._canonical_hash(
        builder._figure_suite_identity_preimage(
            status="final",
            figure_hashes=figure_hashes,
            source_hashes=sources,
            figure_input_assembly_identity_sha256=payload[
                "figure_input_assembly_identity_sha256"
            ],
            model_contract_proposal_identity_sha256=payload[
                "model_contract_proposal_identity_sha256"
            ],
            model_contract_public_identity=public_identity,
            train399_prediction_input_provenance=provenance,
            supplementary_table_bundle_identity_sha256=(
                supplementary_table_bundle_identity
            ),
            supplementary_table_bundle_receipt_sha256=(
                supplementary_table_bundle_receipt
            ),
        )
    )
    return _write_json(root / "summary.json", payload)


def _final_fixture(tmp_path: Path):
    tmp_path.mkdir(parents=True, exist_ok=True)
    manuscript = _manuscript(tmp_path / "compiled.md", placeholders=False)
    metadata = _metadata(tmp_path / "metadata.json", final=True)
    compile_receipt_path = tmp_path / "compile.json"
    compile_receipt = _compile_receipt(compile_receipt_path, manuscript)
    figure_summary = _figure_summary(tmp_path / "figures", compile_receipt)
    return manuscript, metadata, compile_receipt_path, figure_summary


def test_final_docx_is_deterministic_hash_closed_and_structurally_complete(
    tmp_path: Path,
) -> None:
    manuscript, metadata, compile_receipt, figure_summary = _final_fixture(tmp_path)
    anonymous_outputs = []
    title_outputs = []
    receipts = []
    for suffix in ("a", "b"):
        output = tmp_path / f"anonymous-main-{suffix}.docx"
        title = tmp_path / f"title-page-{suffix}.docx"
        receipt = tmp_path / f"submission-{suffix}.json"
        result = builder.build_submission_docx(
            mode="final",
            manuscript=manuscript,
            submission_metadata=metadata,
            compile_receipt=compile_receipt,
            figure_summary=figure_summary,
            title_page_output=title,
            anonymized_main_output=output,
            receipt=receipt,
        )
        assert result["status"] == "completed_final_double_anonymous_submission_bundle"
        assert result["submission_use_allowed"] is True
        assert result["main_figure_count"] == 6
        assert result["main_table_count"] == 3
        assert result["main_reference_count"] == 2
        assert result["abstract_word_count"] == 10
        assert result["abstract_word_limit"] == 249
        assert result["docx_sha256"] == builder._sha256_file(output)
        anonymous_outputs.append(output)
        title_outputs.append(title)
        receipts.append(receipt)
    assert anonymous_outputs[0].read_bytes() == anonymous_outputs[1].read_bytes()
    assert title_outputs[0].read_bytes() == title_outputs[1].read_bytes()
    assert receipts[0].read_bytes() == receipts[1].read_bytes()

    document = Document(anonymous_outputs[0])
    assert len(document.tables) == 3
    assert len(document.inline_shapes) == 6
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Machine-Fill Placeholder Registry" not in text
    assert "[PLACE FIGURE" not in text
    assert "{{" not in text
    assert "Verified Test Author" not in text
    assert "Verified Test Institute" not in text
    assert "verified@example.org" not in text
    assert "Acknowledgments" not in text
    assert "Funding" not in text
    assert "Author Contributions" not in text
    assert "Competing Interests" not in text
    assert "Data Availability" in text
    assert "Code Availability" in text
    assert "https://example.org/phaxis-data" in text
    assert "https://example.org/phaxis-code" in text
    title_text = "\n".join(
        paragraph.text for paragraph in Document(title_outputs[0]).paragraphs
    )
    assert "EDITOR-ONLY TITLE PAGE" in title_text
    assert "Verified Test Author" in title_text
    assert "Verified Test Institute" in title_text
    assert "1 Verified Way" in title_text
    assert "Verified Grant 123" in title_text
    title_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "PHAxis submission fixture"
    )
    assert title_paragraph.style.name != "Title"
    assert document.sections[0].different_first_page_header_footer is False
    assert all(
        section.different_first_page_header_footer is False
        for section in document.sections[1:]
    )

    with zipfile.ZipFile(anonymous_outputs[0]) as archive:
        document_xml = archive.read("word/document.xml")
        numbering_xml = archive.read("word/numbering.xml")
    assert b"w:lnNumType" in document_xml
    assert b"w:insideV w:val=\"nil\"" in document_xml
    assert b"w:tblLayout w:type=\"fixed\"" in document_xml
    assert b"w:numFmt w:val=\"decimal\"" in numbering_xml
    assert b"w:tab w:val=\"num\" w:pos=\"720\"" in numbering_xml
    assert b"w:ind w:left=\"720\" w:hanging=\"360\"" in numbering_xml
    assert document_xml.count(b"w:suppressLineNumbers") >= 3
    assert b"w:pBdr" not in document_xml

    for table in document.tables:
        tbl_pr = table._tbl.tblPr
        total = int(tbl_pr.first_child_found_in("w:tblW").get(qn("w:w")))
        indent = int(tbl_pr.first_child_found_in("w:tblInd").get(qn("w:w")))
        widths = [int(column.get(qn("w:w"))) for column in table._tbl.tblGrid]
        assert indent == 120
        assert sum(widths) == total
        for row in table.rows:
            assert [int(cell._tc.tcPr.tcW.get(qn("w:w"))) for cell in row.cells] == widths
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = tr_pr.find(qn("w:cantSplit"))
            assert cant_split is not None
            assert cant_split.get(qn("w:val")) == "true"
        assert all(
            paragraph.paragraph_format.keep_with_next is True
            for cell in table.rows[0].cells
            for paragraph in cell.paragraphs
        )


def test_layout_fixture_is_visibly_non_submission_and_allows_machine_tokens(
    tmp_path: Path,
) -> None:
    manuscript = _manuscript(tmp_path / "master.md", placeholders=True)
    metadata = _metadata(tmp_path / "metadata.json", final=False)
    output = tmp_path / "layout.docx"
    title = tmp_path / "layout-title.docx"
    result = builder.build_submission_docx(
        mode="layout-fixture",
        manuscript=manuscript,
        submission_metadata=metadata,
        title_page_output=title,
        anonymized_main_output=output,
    )
    assert result["submission_use_allowed"] is False
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "LAYOUT FIXTURE - NOT FOR SUBMISSION" in text
    assert "{{FINAL_HAIR_F1_20UM}}" in text
    assert text.count("FIGURE ") >= 6


def test_layout_fixture_supports_long_journal_facing_filename(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    manuscript = _manuscript(tmp_path / "master.md", placeholders=True)
    metadata = _metadata(tmp_path / "metadata.json", final=False)
    output = tmp_path / (
        "PHAxis_Plant_Phenomics_Main_Submission_"
        "Six_Figures_Three_Tables_LAYOUT_FIXTURE.docx"
    )
    title = tmp_path / "fixture-title.docx"
    staged_names: list[str] = []
    publish = builder._publish_no_overwrite

    def observe_staging_name(source: Path, destination: Path) -> None:
        staged_names.append(source.name)
        publish(source, destination)

    monkeypatch.setattr(builder, "_publish_no_overwrite", observe_staging_name)
    result = builder.build_submission_docx(
        mode="layout-fixture",
        manuscript=manuscript,
        submission_metadata=metadata,
        title_page_output=title,
        anonymized_main_output=output,
    )
    assert output.is_file()
    assert staged_names == ["title.docx", "anonymous.docx", "receipt.json"]
    assert result["docx_sha256"] == builder._sha256_file(output)


def test_final_rejects_placeholder_or_figure_byte_drift(tmp_path: Path) -> None:
    manuscript, metadata, compile_receipt, figure_summary = _final_fixture(tmp_path)
    manuscript.write_text(
        manuscript.read_text(encoding="utf-8").replace("0.912", "{{FINAL_HAIR_F1_20UM}}"),
        encoding="utf-8",
    )
    with pytest.raises(builder.SubmissionDocxError, match="compile receipt/manuscript SHA mismatch"):
        builder.build_submission_docx(
            mode="final",
            manuscript=manuscript,
            submission_metadata=metadata,
            compile_receipt=compile_receipt,
            figure_summary=figure_summary,
            output=tmp_path / "blocked-placeholder.docx",
        )
    assert not (tmp_path / "blocked-placeholder.docx").exists()

    manuscript, metadata, compile_receipt, figure_summary = _final_fixture(tmp_path / "drift")
    summary = json.loads(figure_summary.read_text(encoding="utf-8"))
    first = Path(summary["figures"][builder.FIGURE_STEMS[0]]["bundle"]["files"]["png"])
    first.write_bytes(first.read_bytes() + b"drift")
    with pytest.raises(builder.SubmissionDocxError, match="figure 1 PNG hash mismatch"):
        builder.build_submission_docx(
            mode="final",
            manuscript=manuscript,
            submission_metadata=metadata,
            compile_receipt=compile_receipt,
            figure_summary=figure_summary,
            output=tmp_path / "blocked-figure.docx",
        )


def test_metadata_template_is_intentionally_incomplete_and_rejected(
    tmp_path: Path,
) -> None:
    manuscript = _manuscript(tmp_path / "master.md", placeholders=True)
    frontmatter = builder._extract_frontmatter(manuscript.read_text(encoding="utf-8"))
    template = builder.submission_metadata_template(frontmatter=frontmatter)
    metadata = _write_json(tmp_path / "template.json", template)
    with pytest.raises(builder.SubmissionDocxError, match="submission metadata status is invalid"):
        builder.build_submission_docx(
            mode="layout-fixture",
            manuscript=manuscript,
            submission_metadata=metadata,
            output=tmp_path / "blocked.docx",
        )


def test_metadata_v2_template_cli_is_create_only_and_contains_postal_gate(
    tmp_path: Path,
) -> None:
    manuscript = _manuscript(tmp_path / "master.md", placeholders=True)
    destination = tmp_path / "metadata-v2.json"
    assert (
        builder.main(
            [
                "--manuscript",
                str(manuscript),
                "--write-metadata-template",
                str(destination),
            ]
        )
        == 0
    )
    original = destination.read_bytes()
    payload = json.loads(original)
    assert payload["schema_version"] == "PHAxis-submission-title-metadata-2.0"
    assert payload["authors"][0]["postal_address"] is None
    assert payload["authors"][0]["postal_address_author_verified"] is False
    unsigned = dict(payload)
    identity = unsigned.pop("metadata_identity_sha256")
    assert builder._canonical_hash(unsigned) == identity
    with pytest.raises(builder.SubmissionDocxError, match="refusing to overwrite"):
        builder.main(
            [
                "--manuscript",
                str(manuscript),
                "--write-metadata-template",
                str(destination),
            ]
        )
    assert destination.read_bytes() == original


def test_checked_in_metadata_v2_is_self_sealed_and_v1_history_is_unchanged() -> None:
    directory = PROJECT_ROOT / "manuscript/phaxis_v1_0"
    historical = directory / "SUBMISSION_TITLE_METADATA_TEMPLATE.json"
    version_two = directory / "SUBMISSION_TITLE_METADATA_TEMPLATE_2_0.json"
    assert builder._sha256_file(historical) == (
        "ce2b911fb6f22cdfbccc3c4290865d4a5167ff4b76af648ba389bcccf5572c6b"
    )
    assert json.loads(historical.read_text(encoding="utf-8"))["schema_version"] == (
        "PHAxis-submission-title-metadata-1.0"
    )
    payload = json.loads(version_two.read_text(encoding="utf-8"))
    assert payload["schema_version"] == builder.METADATA_SCHEMA
    assert payload["authors"][0]["postal_address"] is None
    assert payload["authors"][0]["postal_address_author_verified"] is False
    identity = payload.pop("metadata_identity_sha256")
    assert builder._canonical_hash(payload) == identity


def test_final_metadata_v2_requires_verified_postal_address_and_native_seal(
    tmp_path: Path,
) -> None:
    manuscript, metadata, compile_receipt, figure_summary = _final_fixture(tmp_path)
    payload = json.loads(metadata.read_text(encoding="utf-8"))
    payload["authors"][0]["postal_address"] = ""
    payload["metadata_identity_sha256"] = builder._canonical_hash(
        {key: value for key, value in payload.items() if key != "metadata_identity_sha256"}
    )
    _write_json(metadata, payload)
    with pytest.raises(builder.SubmissionDocxError, match="postal address is empty"):
        builder.build_submission_docx(
            mode="final",
            manuscript=manuscript,
            submission_metadata=metadata,
            compile_receipt=compile_receipt,
            figure_summary=figure_summary,
            title_page_output=tmp_path / "title.docx",
            anonymized_main_output=tmp_path / "anonymous.docx",
        )

    payload["authors"][0]["postal_address"] = "1 Verified Way"
    # Deliberately retain the seal for the empty address.
    _write_json(metadata, payload)
    with pytest.raises(builder.SubmissionDocxError, match="identity seal mismatch"):
        builder.build_submission_docx(
            mode="final",
            manuscript=manuscript,
            submission_metadata=metadata,
            compile_receipt=compile_receipt,
            figure_summary=figure_summary,
            title_page_output=tmp_path / "title.docx",
            anonymized_main_output=tmp_path / "anonymous.docx",
        )


def test_three_member_publication_rolls_back_and_never_overwrites(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    manuscript, metadata, compile_receipt, figure_summary = _final_fixture(tmp_path)
    title = tmp_path / "title.docx"
    anonymous = tmp_path / "anonymous.docx"
    receipt = tmp_path / "receipt.json"
    publish = builder._publish_no_overwrite
    calls = 0

    def fail_second(source: Path, destination: Path) -> None:
        nonlocal calls
        calls += 1
        if calls == 2:
            raise OSError("injected second-member publication failure")
        publish(source, destination)

    monkeypatch.setattr(builder, "_publish_no_overwrite", fail_second)
    with pytest.raises(OSError, match="injected"):
        builder.build_submission_docx(
            mode="final",
            manuscript=manuscript,
            submission_metadata=metadata,
            compile_receipt=compile_receipt,
            figure_summary=figure_summary,
            title_page_output=title,
            anonymized_main_output=anonymous,
            receipt=receipt,
        )
    assert not title.exists() and not anonymous.exists() and not receipt.exists()

    monkeypatch.setattr(builder, "_publish_no_overwrite", publish)
    title.write_bytes(b"preexisting")
    with pytest.raises(builder.SubmissionDocxError, match="overwrite"):
        builder.build_submission_docx(
            mode="final",
            manuscript=manuscript,
            submission_metadata=metadata,
            compile_receipt=compile_receipt,
            figure_summary=figure_summary,
            title_page_output=title,
            anonymized_main_output=anonymous,
            receipt=receipt,
        )
    assert title.read_bytes() == b"preexisting"
    assert not anonymous.exists() and not receipt.exists()
