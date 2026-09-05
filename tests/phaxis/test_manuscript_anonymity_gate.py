from __future__ import annotations

import importlib.util
import io
from pathlib import Path
import zipfile

from docx import Document
from docx.shared import Inches
from PIL import Image, PngImagePlugin
import pytest


PROJECT_ROOT = Path(__file__).resolve().parents[2]
SCRIPT = PROJECT_ROOT / "scripts/phaxis/verify_manuscript_artifacts.py"
SPEC = importlib.util.spec_from_file_location("phaxis_anonymity_gate", SCRIPT)
assert SPEC and SPEC.loader
gate = importlib.util.module_from_spec(SPEC)
SPEC.loader.exec_module(gate)


METADATA = {
    "authors": [
        {
            "full_name": "Verified Test Author",
            "email": "verified@example.org",
            "orcid": "0000-0002-1825-0097",
            "postal_address": "1 Verified Way, Test City 10000, Testland",
            "postal_address_author_verified": True,
            "corresponding_author": True,
        }
    ],
    "affiliations": [{"id": 1, "text": "Verified Test Institute"}],
}
DECLARATIONS = {
    "Acknowledgments": "We thank the Verified Imaging Core.",
    "Funding": "This work was supported by Verified Grant 123.",
    "Author Contributions": "Verified Test Author designed the study.",
    "Competing Interests": "The author declares no competing interests.",
}


def test_data_and_code_availability_contract_is_executable_and_fail_closed() -> None:
    good = (
        "## 5. Data Availability\n\n"
        "The dataset card and annotation schema, example input, and expected output "
        "are at https://example.org/data under the CC-BY-4.0 license.\n\n"
        "## 6. Code Availability\n\n"
        "Source, documentation, environment locks, example input, and expected output "
        "are at https://example.org/code under the Apache-2.0 license.\n"
    )
    result = gate._availability_contract(good)
    assert result["data_url_count"] == 1
    assert result["code_url_count"] == 1
    assert result["available_upon_request_only"] is False
    with pytest.raises(gate.ManuscriptArtifactError, match="available-upon-request"):
        gate._availability_contract(
            good.replace(
                "The dataset card and annotation schema, example input, and expected output "
                "are at https://example.org/data under the CC-BY-4.0 license.",
                "Data are available upon reasonable request.",
            )
        )


def _save_document(path: Path, paragraphs: list[str], *, image: Path | None = None) -> Path:
    document = Document()
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    for value in paragraphs:
        document.add_paragraph(value)
    if image is not None:
        document.add_picture(str(image), width=Inches(1.0))
    document.save(path)
    return path


def _rewrite_member(source: Path, destination: Path, name: str, value: bytes) -> None:
    with zipfile.ZipFile(source) as archive:
        members = {info.filename: archive.read(info.filename) for info in archive.infolist()}
    members[name] = value
    with zipfile.ZipFile(destination, "w") as archive:
        for member, raw in sorted(members.items()):
            archive.writestr(member, raw)


def test_title_positive_completeness_and_reviewer_package_pass(tmp_path: Path) -> None:
    denylist = gate._identity_denylist(METADATA, DECLARATIONS)
    title = _save_document(
        tmp_path / "PHAxis_title_page.docx",
        [
            "Verified Test Author",
            "verified@example.org",
            "0000-0002-1825-0097",
            "1 Verified Way, Test City 10000, Testland",
            "Verified Test Institute",
            *[item for pair in DECLARATIONS.items() for item in pair],
        ],
    )
    title_record = gate._inspect_docx(
        title,
        expected_media=0,
        role="title",
        reviewer_visible=False,
        denylist=denylist,
    )
    gate._validate_title_page_completeness(
        title_record,
        metadata=METADATA,
        declarations=DECLARATIONS,
    )

    reviewer = _save_document(
        tmp_path / "PHAxis_anonymized_main.docx",
        [
            "PHAxis anonymous review copy",
            "Root-hair spatial phenotypes were measured.",
            "Funding availability can shape imaging throughput across laboratories.",
        ],
    )
    record = gate._inspect_docx(
        reviewer,
        expected_media=0,
        role="reviewer main",
        reviewer_visible=True,
        denylist=denylist,
    )
    assert record["identity_denylist_hit_count"] == 0
    assert record["tracked_change_count"] == 0
    assert record["hidden_text_count"] == 0


def test_title_completeness_requires_visible_identity_not_only_core_metadata(
    tmp_path: Path,
) -> None:
    title = _save_document(
        tmp_path / "title-missing-visible-author.docx",
        [
            "verified@example.org",
            "0000-0002-1825-0097",
            "1 Verified Way, Test City 10000, Testland",
            "Verified Test Institute",
            *[item for pair in DECLARATIONS.items() for item in pair],
        ],
    )
    document = Document(title)
    document.core_properties.author = "Verified Test Author"
    document.save(title)
    record = gate._inspect_docx(
        title,
        expected_media=0,
        role="incomplete title",
        reviewer_visible=False,
        denylist=gate._identity_denylist(METADATA, DECLARATIONS),
    )
    with pytest.raises(gate.ManuscriptArtifactError, match="author_1_full_name"):
        gate._validate_title_page_completeness(
            record,
            metadata=METADATA,
            declarations=DECLARATIONS,
        )


def test_reviewer_exact_declaration_heading_fails_without_generic_word_false_positive(
    tmp_path: Path,
) -> None:
    reviewer = _save_document(
        tmp_path / "reviewer-declaration-heading.docx",
        ["Anonymous review copy", "Funding"],
    )
    with pytest.raises(
        gate.ManuscriptArtifactError,
        match="editor-only declaration headings",
    ):
        gate._inspect_docx(
            reviewer,
            expected_media=0,
            role="reviewer declaration mutation",
            reviewer_visible=True,
            denylist=gate._identity_denylist(METADATA, DECLARATIONS),
        )


@pytest.mark.parametrize(
    "injection",
    (
        "core",
        "header",
        "footnote",
        "comments",
        "custom_properties",
        "tracked_change",
        "hidden",
        "relationship_identity_slug",
        "relationship",
    ),
)
def test_reviewer_ooxml_identity_mutations_fail_closed(
    tmp_path: Path, injection: str
) -> None:
    denylist = gate._identity_denylist(METADATA, DECLARATIONS)
    clean = _save_document(tmp_path / "clean.docx", ["Anonymous review copy"])
    mutated = tmp_path / f"mutated-{injection}.docx"
    if injection == "core":
        with zipfile.ZipFile(clean) as archive:
            raw = archive.read("docProps/core.xml")
        raw = raw.replace(b"<dc:creator></dc:creator>", b"<dc:creator>Verified Test Author</dc:creator>")
        raw = raw.replace(b"<dc:creator/>", b"<dc:creator>Verified Test Author</dc:creator>")
        _rewrite_member(clean, mutated, "docProps/core.xml", raw)
        message = "identity terms|core creator"
    elif injection == "header":
        document = Document(clean)
        document.sections[0].header.paragraphs[0].text = "Verified Test Author"
        document.save(mutated)
        message = "identity terms"
    elif injection == "footnote":
        raw = (
            f'<?xml version="1.0" encoding="UTF-8"?><w:footnotes xmlns:w="{gate.W_NS}">'
            '<w:footnote w:id="1"><w:p><w:r><w:t>Verified Test Author</w:t>'
            "</w:r></w:p></w:footnote></w:footnotes>"
        ).encode("utf-8")
        _rewrite_member(clean, mutated, "word/footnotes.xml", raw)
        message = "identity terms"
    elif injection == "comments":
        raw = (
            f'<?xml version="1.0" encoding="UTF-8"?><w:comments xmlns:w="{gate.W_NS}"/>'
        ).encode("utf-8")
        _rewrite_member(clean, mutated, "word/comments.xml", raw)
        message = "identity-bearing OOXML parts"
    elif injection == "custom_properties":
        raw = (
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties"/>'
        ).encode("utf-8")
        _rewrite_member(clean, mutated, "docProps/custom.xml", raw)
        message = "identity-bearing OOXML parts"
    elif injection == "tracked_change":
        with zipfile.ZipFile(clean) as archive:
            raw = archive.read("word/document.xml")
        raw = raw.replace(
            b"</w:body>", b'<w:ins w:id="1"/></w:body>', 1
        )
        _rewrite_member(clean, mutated, "word/document.xml", raw)
        message = "tracked changes"
    elif injection == "hidden":
        with zipfile.ZipFile(clean) as archive:
            raw = archive.read("word/document.xml")
        raw = raw.replace(b"<w:t>", b"<w:rPr><w:vanish/></w:rPr><w:t>", 1)
        _rewrite_member(clean, mutated, "word/document.xml", raw)
        message = "hidden text"
    elif injection == "relationship_identity_slug":
        with zipfile.ZipFile(clean) as archive:
            raw = archive.read("word/_rels/document.xml.rels")
        raw = raw.replace(
            b"</Relationships>",
            b'<Relationship Id="rIdLeak" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" Target="https://example.org/profiles/verified-test-author" TargetMode="External"/></Relationships>',
        )
        _rewrite_member(clean, mutated, "word/_rels/document.xml.rels", raw)
        message = "identity terms"
    else:
        with zipfile.ZipFile(clean) as archive:
            raw = archive.read("word/_rels/document.xml.rels")
        raw = raw.replace(
            b"</Relationships>",
            b'<Relationship Id="rIdLeak" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" Target="mailto:verified@example.org" TargetMode="External"/></Relationships>',
        )
        _rewrite_member(clean, mutated, "word/_rels/document.xml.rels", raw)
        message = "external relationship"
    with pytest.raises(gate.ManuscriptArtifactError, match=message):
        gate._inspect_docx(
            mutated,
            expected_media=0,
            role="mutated reviewer",
            reviewer_visible=True,
            denylist=denylist,
        )


def test_reviewer_filename_identity_injection_fails_closed(tmp_path: Path) -> None:
    reviewer = _save_document(
        tmp_path / "Verified Test Author reviewer copy.docx",
        ["Anonymous review copy"],
    )
    with pytest.raises(gate.ManuscriptArtifactError, match="identity terms"):
        gate._inspect_docx(
            reviewer,
            expected_media=0,
            role="reviewer filename mutation",
            reviewer_visible=True,
            denylist=gate._identity_denylist(METADATA, DECLARATIONS),
        )


def test_embedded_png_identity_metadata_fails_closed(tmp_path: Path) -> None:
    image_path = tmp_path / "identity.png"
    info = PngImagePlugin.PngInfo()
    info.add_text("Author", "Verified Test Author")
    Image.new("RGB", (64, 64), "white").save(image_path, pnginfo=info)
    reviewer = _save_document(
        tmp_path / "anonymous-with-image.docx",
        ["Anonymous review copy"],
        image=image_path,
    )
    with pytest.raises(gate.ManuscriptArtifactError, match="embedded image metadata"):
        gate._inspect_docx(
            reviewer,
            expected_media=1,
            role="reviewer image",
            reviewer_visible=True,
            denylist=gate._identity_denylist(METADATA, DECLARATIONS),
        )


def test_stage55_two_member_publication_is_all_or_none(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    first = tmp_path / "first.tmp"
    second = tmp_path / "second.tmp"
    first.write_bytes(b"first")
    second.write_bytes(b"second")
    first_output = tmp_path / "upload.json"
    second_output = tmp_path / "receipt.json"
    publish = gate._publish_create_only
    calls = 0

    def fail_second(source: Path, destination: Path) -> None:
        nonlocal calls
        calls += 1
        if calls == 2:
            raise OSError("injected stage55 second-member failure")
        publish(source, destination)

    monkeypatch.setattr(gate, "_publish_create_only", fail_second)
    with pytest.raises(OSError, match="injected"):
        gate._publish_transaction(
            ((first, first_output), (second, second_output))
        )
    assert not first_output.exists() and not second_output.exists()
