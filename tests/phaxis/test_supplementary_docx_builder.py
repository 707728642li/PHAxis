from __future__ import annotations

import hashlib
import importlib.util
import inspect
import json
from pathlib import Path
import zipfile

from docx import Document
from docx.oxml.ns import qn
from PIL import Image
import pytest

from phaxis.supplementary_tables import (
    BUNDLE_RECEIPT as SUPPLEMENTARY_TABLE_BUNDLE_RECEIPT,
    BUNDLE_DIRECTORY as SUPPLEMENTARY_TABLE_BUNDLE_DIRECTORY,
    FINAL_STATUS as FINAL_SUPPLEMENTARY_TABLE_STATUS,
    materialize_supplementary_table_data_bundle,
)
from tests.phaxis.test_supplementary_table_data_bundle import (
    source_fixture as _supplementary_source_fixture,
)


PROJECT_ROOT = Path(__file__).resolve().parents[2]
SCRIPT = PROJECT_ROOT / "scripts/phaxis/build_supplementary_docx.py"
SPEC = importlib.util.spec_from_file_location("phaxis_supplementary_docx", SCRIPT)
assert SPEC and SPEC.loader
builder = importlib.util.module_from_spec(SPEC)
SPEC.loader.exec_module(builder)


def _write_json(path: Path, payload: dict) -> Path:
    path.write_text(
        json.dumps(payload, ensure_ascii=False, allow_nan=False, indent=2, sort_keys=True)
        + "\n",
        encoding="utf-8",
    )
    return path


def _supplement(path: Path, *, final: bool) -> Path:
    status = (
        "final sealed supplement bound to the PHAxis 1.0.0 evidence graph."
        if final
        else "evidence-bound master; final rendering requires sealed evidence."
    )
    methods = []
    for number in range(1, 10):
        methods.append(
            f"### S{number}. Supplementary method {number}\n\n"
            f"Method {number} preserves biological interpretation and physical units.\n"
        )
    figure_legends = []
    for number, _, title in builder.SUPPLEMENTARY_FIGURES:
        figure_legends.append(
            f"### Figure S{number}. {title}\n\n"
            "Panels report the declared source-unit measurement and observability.\n"
        )
    table_legends = []
    for number in range(1, 11):
        table_legends.append(
            f"### Table S{number}. Supplementary data resource {number}\n\n"
            "The resource retains units, denominators, and provenance.\n"
        )
    path.write_text(
        "# Supplementary Materials for PHAxis submission fixture\n\n"
        "**Software and model-system version:** PHAxis 1.0.0  \n"
        "**Companion main manuscript:** `main.md`  \n"
        f"**Status:** {status}\n\n"
        "## Supplementary Methods\n\n"
        + "\n".join(methods)
        + "\n### S2a. Physical input channels\n\n"
        "1. robust normalized intensity;\n"
        "2. local physical contrast;\n"
        "3. dark-ridge response.\n\n"
        "## Supplementary Figure Legends\n\n"
        + "\n".join(figure_legends)
        + "\n## Supplementary Tables and Data Files\n\n"
        + "\n".join(table_legends)
        + "\n## Supplementary machine-fill policy\n\n"
        "Final values come from the same sealed evidence graph as the main manuscript.\n",
        encoding="utf-8",
    )
    return path


def _main_manuscript(path: Path) -> Path:
    path.write_text(
        "# PHAxis submission fixture\n\n"
        "**Running title:** PHAxis fixture  \n"
        "**Article type:** Software and Hardware Article  \n"
        "**Software version:** PHAxis 1.0.0  \n\n"
        "## Abstract\n\nFinal compiled main manuscript.\n",
        encoding="utf-8",
    )
    return path


def _metadata(path: Path) -> Path:
    payload = {
        "schema_version": builder.main.METADATA_SCHEMA,
        "status": builder.main.FINAL_METADATA_STATUS,
        "journal": "Plant Phenomics",
        "manuscript_title": "PHAxis submission fixture",
        "running_title": "PHAxis fixture",
        "article_type": "Software and Hardware Article",
        "software_version": "PHAxis 1.0.0",
        "authors": [
            {
                "full_name": "Verified Test Author",
                "affiliation_ids": [1],
                "corresponding_author": True,
                "email": "verified@example.org",
                "orcid": "0000-0002-1825-0097",
                "postal_address": "1 Verified Way, Test City 10000, Testland",
                "postal_address_author_verified": True,
            }
        ],
        "affiliations": [{"id": 1, "text": "Verified Test Institute"}],
    }
    payload["metadata_identity_sha256"] = builder.main._canonical_hash(payload)
    return _write_json(path, payload)


def _compile_receipt(path: Path, manuscript: Path) -> Path:
    manuscript_text = manuscript.read_text(encoding="utf-8")
    abstract_words = builder.main.require_abstract_within_limit(manuscript_text)
    payload = {
        "schema_version": builder.main.COMPILE_RECEIPT_SCHEMA,
        "status": builder.main.FINAL_COMPILE_STATUS,
        "output_sha256": builder.main._sha256_file(manuscript),
        "unresolved_token_count": 0,
        "author_metadata_complete": True,
        "abstract_word_count": abstract_words,
        "abstract_word_limit": builder.main.ABSTRACT_WORD_LIMIT,
        "abstract_word_limit_passed": True,
        "blind_images_used": 0,
        "root_cap_region_statistics_included": False,
        "model_bundle_id": "PHAXIS-V1.0.0-STRICT-TRAIN399-TEST",
        "root_expert_id": "PHAxis-root-provider-TEST",
        "hair_identity_count_expert": "PHAxis-StageB-train399-five-seed",
        "model_contract_proposal_identity_sha256": hashlib.sha256(
            b"model-contract-proposal"
        ).hexdigest(),
        "figure_input_assembly_identity_sha256": hashlib.sha256(
            b"figure-input-assembly"
        ).hexdigest(),
    }
    payload["receipt_identity_sha256"] = builder.main._canonical_hash(payload)
    return _write_json(path, payload)


def _figure_suite(path: Path, compile_receipt: Path) -> Path:
    path.mkdir(parents=True)
    compile_payload = json.loads(compile_receipt.read_text(encoding="utf-8"))
    supplementary_sources = _supplementary_source_fixture(
        path / "supplementary-table-authorities"
    )
    supplementary_table_bundle = materialize_supplementary_table_data_bundle(
        output=path / SUPPLEMENTARY_TABLE_BUNDLE_DIRECTORY,
        status=FINAL_SUPPLEMENTARY_TABLE_STATUS,
        source_paths=supplementary_sources,
        source_identities={},
        figure_input_manifest_sha256=hashlib.sha256(
            b"figure-input-manifest"
        ).hexdigest(),
        figure_input_assembly_identity_sha256=compile_payload[
            "figure_input_assembly_identity_sha256"
        ],
        model_contract_proposal_identity_sha256=compile_payload[
            "model_contract_proposal_identity_sha256"
        ],
    )
    main_figures = {}
    main_hashes = {}
    for number, stem in enumerate(builder.main.FIGURE_STEMS, start=1):
        image_path = path / f"{stem}.png"
        Image.new("RGB", (1200, 700), (245, 247, 250)).save(image_path)
        main_figures[stem] = {
            "number": number,
            "title": stem.replace("_", " "),
            "status": "final",
            "bundle": {"files": {"png": str(image_path.resolve())}},
        }
        main_hashes[stem] = {
            "png": builder.main._sha256_file(image_path),
            "pdf": hashlib.sha256(f"main-pdf-{number}".encode()).hexdigest(),
            "tiff": hashlib.sha256(f"main-tiff-{number}".encode()).hexdigest(),
            "source_data": {},
        }
    supplementary_figures = {}
    supplementary_hashes = {}
    for number, stem, title in builder.SUPPLEMENTARY_FIGURES:
        image_path = path / f"{stem}.png"
        Image.new("RGB", (1200, 700), (245, 247, 250)).save(image_path)
        supplementary_figures[stem] = {
            "number": f"S{number}",
            "title": title,
            "status": "final",
            "bundle": {"files": {"png": str(image_path.resolve())}},
        }
        supplementary_hashes[stem] = {
            "png": builder.main._sha256_file(image_path),
            "pdf": hashlib.sha256(f"supp-pdf-{number}".encode()).hexdigest(),
            "tiff": hashlib.sha256(f"supp-tiff-{number}".encode()).hexdigest(),
            "source_data": {},
        }
    public_identity = {
        "model_bundle_id": compile_payload["model_bundle_id"],
        "root_expert_id": compile_payload["root_expert_id"],
    }
    sources = {"stageb": hashlib.sha256(b"stageb").hexdigest()}
    provenance = {"ordered_file_set_identity_sha256": hashlib.sha256(b"ordered").hexdigest()}
    payload = {
        "schema_version": builder.FIGURE_SUITE_SCHEMA,
        "status": builder.FINAL_FIGURE_STATUS,
        "submission_use_allowed": True,
        "formal_train399_only_gate_passed": True,
        "blind_images_used": 0,
        "claim_contract": {
            "main_figure_count": 6,
            "supplementary_figure_count": 9,
            "supplementary_table_data_resource_count": 10,
            "root_cap_region_statistics_included": False,
            "canonical_annotations_read": False,
        },
        "model_contract_proposal_identity_sha256": compile_payload.get("model_contract_proposal_identity_sha256"),
        "figure_input_assembly_identity_sha256": compile_payload.get("figure_input_assembly_identity_sha256"),
        "model_bundle_id": compile_payload["model_bundle_id"],
        "root_expert_id": compile_payload["root_expert_id"],
        "hair_identity_expert_id": compile_payload[
            "hair_identity_count_expert"
        ],
        "model_contract_public_identity": public_identity,
        "train399_prediction_input_provenance": provenance,
        "source_summary_sha256": sources,
        "figure_bundle_sha256": main_hashes,
        "figures": main_figures,
        "supplementary_figures": supplementary_figures,
        "supplementary_figure_bundle_sha256": supplementary_hashes,
        "supplementary_figure_bundle_identity_sha256": builder.main._canonical_hash(supplementary_hashes),
        "supplementary_tables": supplementary_table_bundle["items"],
        "supplementary_table_bundle_receipt": (
            f"{SUPPLEMENTARY_TABLE_BUNDLE_DIRECTORY}/"
            f"{SUPPLEMENTARY_TABLE_BUNDLE_RECEIPT}"
        ),
        "supplementary_table_bundle_receipt_sha256": (
            supplementary_table_bundle["receipt_sha256"]
        ),
        "supplementary_table_bundle_identity_sha256": (
            supplementary_table_bundle["bundle_identity_sha256"]
        ),
        "supplementary_table_bundle_sha256": supplementary_table_bundle[
            "bundle_file_sha256"
        ],
    }
    payload["figure_suite_identity_sha256"] = builder.main._canonical_hash(
        builder.main._figure_suite_identity_preimage(
            status="final",
            figure_hashes=main_hashes,
            source_hashes=sources,
            figure_input_assembly_identity_sha256=payload[
                "figure_input_assembly_identity_sha256"
            ],
            model_contract_proposal_identity_sha256=payload[
                "model_contract_proposal_identity_sha256"
            ],
            model_contract_public_identity=public_identity,
            train399_prediction_input_provenance=provenance,
            supplementary_table_bundle_identity_sha256=(
                supplementary_table_bundle["bundle_identity_sha256"]
            ),
            supplementary_table_bundle_receipt_sha256=(
                supplementary_table_bundle["receipt_sha256"]
            ),
        )
    )
    return _write_json(path / "summary.json", payload)


def _final_fixture(tmp_path: Path):
    tmp_path.mkdir(parents=True, exist_ok=True)
    supplement = _supplement(tmp_path / "supplement.md", final=True)
    main_manuscript = _main_manuscript(tmp_path / "main.md")
    metadata = _metadata(tmp_path / "metadata.json")
    compile_receipt = _compile_receipt(tmp_path / "compile.json", main_manuscript)
    figures = _figure_suite(tmp_path / "figures", compile_receipt)
    return supplement, main_manuscript, metadata, compile_receipt, figures


def test_layout_fixture_is_deterministic_visible_and_structurally_complete(
    tmp_path: Path,
) -> None:
    supplement = _supplement(tmp_path / "supplement.md", final=False)
    outputs = []
    for suffix in ("a", "b"):
        output = tmp_path / f"supplement-{suffix}.docx"
        result = builder.build_supplementary_docx(
            mode="layout-fixture",
            supplement=supplement,
            output=output,
        )
        assert result["submission_use_allowed"] is False
        assert result["supplementary_figure_count"] == 9
        assert result["supplementary_table_data_resource_count"] == 10
        outputs.append(output)
    assert outputs[0].read_bytes() == outputs[1].read_bytes()

    document = Document(outputs[0])
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "LAYOUT FIXTURE - NOT FOR SUBMISSION" in text
    assert text.count("LAYOUT PLACEHOLDER") == 9
    assert "Supplementary method 9" in text
    assert "Table S10" in text
    assert "Supplementary machine-fill policy" not in text
    assert "Final values come from the same sealed evidence graph" not in text
    assert len(document.inline_shapes) == 0
    assert document.sections[0].different_first_page_header_footer is True
    assert all(
        section.different_first_page_header_footer is False
        for section in document.sections[1:]
    )
    with zipfile.ZipFile(outputs[0]) as archive:
        document_xml = archive.read("word/document.xml")
        numbering_xml = archive.read("word/numbering.xml")
    assert b"w:lnNumType" in document_xml
    assert b"w:numFmt w:val=\"decimal\"" in numbering_xml
    assert b"w:pBdr" not in document_xml


def test_real_master_layout_fixture_embeds_table_s9_as_four_column_table(
    tmp_path: Path,
) -> None:
    supplement = (
        PROJECT_ROOT
        / "docs/phaxis/PHAXIS_SUPPLEMENTARY_MASTER_DRAFT_20260830.md"
    )
    output = tmp_path / "real-master-layout.docx"
    supplement_text = supplement.read_text(encoding="utf-8")
    assert "Nine tensor heads predict" in supplement_text
    assert "all nine tensor heads" in supplement_text
    assert "base direction and base length jointly encode" in supplement_text
    assert "Expanded acquisition-challenge overlay gallery" in supplement_text
    assert "RHSCU-aa5b6e37df15821f" in supplement_text
    assert "RHSCU-bbf649822174e0a2" in supplement_text
    assert "condition-blinded overlay gallery" not in supplement_text.casefold()
    result = builder.build_supplementary_docx(
        mode="layout-fixture",
        supplement=supplement,
        output=output,
    )

    expected_rows = [
        [
            "S9 block",
            "Row unit",
            "Mandatory biological/statistical content",
            "Interpretation",
        ],
        [
            "A — raw condition atlas",
            "One trait × one cohort × one D15 condition",
            "Trait ID/field/measurement family/unit; total n; non-null n; "
            "observability; median; Q25; Q75; IQR; range; status/reason",
            "Raw source-unit phenotype distribution and measurement "
            "availability before model fitting",
        ],
        [
            "B — effect-status ledger",
            "One trait × one cohort × one contrast",
            "Status; estimate; 95% interval; effect scale; endpoint n; "
            "not_estimable_reason",
            "Complete prespecified-analysis ledger; unestimated is not zero effect",
        ],
        [
            "C — cohort provenance",
            "One source unit or declared denominator cell",
            "clean/full membership; overlap identity; formal/review state; "
            "endpoint support and censoring",
            "Audit trail linking Table S9A/S9B to the application cohort",
        ],
        [
            "D — wt_gate_flow",
            "One cohort × WT experiment × endpoint",
            "Day status; 22°C/30°C total, formal, and endpoint n; "
            "base/endpoint gates; model status/reason",
            "Experiment inventory and transparent estimability flow, with no "
            "phenotype outlier removal",
        ],
        [
            "E — wt_experiment_contrasts",
            "One cohort × WT experiment × endpoint",
            "Within-experiment 30:22°C ratio/interval or typed failure; BH "
            "family; day/meta eligibility",
            "Experiment-blocked WT association; unknown-day contrasts remain "
            "experiment-level only",
        ],
        [
            "F — wt_same_day_meta",
            "One cohort × known day × endpoint",
            "Eligible experiment set and k; REML/Hartung–Knapp ratio/interval; "
            "BH family; τ²/Q/I² or typed failure",
            "Same-day replication summary only; no cross-day, unknown-day, or "
            "clean/full pooling",
        ],
    ]

    document = Document(output)
    rendered_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Supplementary machine-fill policy" not in rendered_text
    # The biology-forward S8 endpoint/use table and the machine-fill S9 ledger
    # must both remain real Word tables; identify S9 by its exact header below.
    assert result["embedded_markdown_table_count"] == 2
    matching_tables = [
        table
        for table in document.tables
        if [cell.text for cell in table.rows[0].cells] == expected_rows[0]
    ]
    assert len(matching_tables) == 1
    table = matching_tables[0]
    assert len(table.columns) == 4
    assert all(len(row.cells) == 4 for row in table.rows)
    assert [[cell.text for cell in row.cells] for row in table.rows] == expected_rows

    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.trPr
        assert tr_pr is not None
        cant_split = tr_pr.find(qn("w:cantSplit"))
        assert cant_split is not None
        assert cant_split.get(qn("w:val")) == "true"
        repeat_header = tr_pr.find(qn("w:tblHeader"))
        if row_index == 0:
            assert repeat_header is not None
            assert repeat_header.get(qn("w:val")) == "true"
            assert all(
                paragraph.paragraph_format.keep_with_next is True
                for cell in row.cells
                for paragraph in cell.paragraphs
            )
        else:
            assert repeat_header is None


def test_layout_fixture_supports_long_journal_facing_filename(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    supplement = _supplement(tmp_path / "supplement.md", final=False)
    output = tmp_path / (
        "PHAxis_Plant_Phenomics_Supplementary_Materials_"
        "Nine_Figures_Ten_Data_Resources_LAYOUT_FIXTURE.docx"
    )
    staged_names: list[str] = []
    publish = builder.main._publish_no_overwrite

    def observe_staging_name(source: Path, destination: Path) -> None:
        staged_names.append(source.name)
        publish(source, destination)

    monkeypatch.setattr(
        builder.main, "_publish_no_overwrite", observe_staging_name
    )
    result = builder.build_supplementary_docx(
        mode="layout-fixture",
        supplement=supplement,
        output=output,
    )
    assert output.is_file()
    assert staged_names == ["document.docx", "receipt.json"]
    assert result["docx_sha256"] == builder.main._sha256_file(output)


def test_final_supplement_is_hash_closed_and_embeds_nine_figures(
    tmp_path: Path,
) -> None:
    supplement, main_manuscript, _metadata_path, compile_receipt, figures = _final_fixture(
        tmp_path
    )
    output = tmp_path / "final-supplement.docx"
    result = builder.build_supplementary_docx(
        mode="final",
        supplement=supplement,
        output=output,
        main_manuscript=main_manuscript,
        main_compile_receipt=compile_receipt,
        figure_summary=figures,
    )
    assert result["submission_use_allowed"] is True
    assert result["status"] == "completed_final_anonymized_supplementary_docx"
    assert result["submission_metadata_consumed"] is False
    assert result["docx_sha256"] == builder.main._sha256_file(output)
    assert result["hair_identity_count_expert"] == "PHAxis-StageB-train399-five-seed"
    document = Document(output)
    assert len(document.inline_shapes) == 9
    assert result["embedded_markdown_table_count"] == len(document.tables)
    file_indexes = [
        table
        for table in document.tables
        if [cell.text for cell in table.rows[0].cells]
        == [
            "Resource",
            "Title",
            "Bundle-relative files",
            "Denominator closure",
            "Item identity (SHA-256)",
        ]
    ]
    assert len(file_indexes) == 1
    index_text = "\n".join(
        cell.text for row in file_indexes[0].rows[1:] for cell in row.cells
    )
    assert index_text.count("item_receipt.json") == 10
    assert "S10/src/s01.csv" in index_text
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "LAYOUT FIXTURE" not in text
    assert "{{" not in text
    assert "Verified Test Author" not in text
    assert "Verified Test Institute" not in text
    with zipfile.ZipFile(output) as archive:
        core = archive.read("docProps/core.xml")
    assert b"Verified Test Author" not in core
    assert b"Verified Test Institute" not in core


def test_final_rejects_provisional_source_and_figure_byte_drift(
    tmp_path: Path,
) -> None:
    supplement, main_manuscript, _metadata_path, compile_receipt, figures = _final_fixture(
        tmp_path
    )
    supplement.write_text(
        supplement.read_text(encoding="utf-8").replace(
            "final sealed supplement bound to the PHAxis 1.0.0 evidence graph.",
            "evidence-bound master; final rendering requires sealed evidence.",
        ),
        encoding="utf-8",
    )
    with pytest.raises(builder.SupplementaryDocxError, match="provisional marker"):
        builder.build_supplementary_docx(
            mode="final",
            supplement=supplement,
            output=tmp_path / "blocked-provisional.docx",
            main_manuscript=main_manuscript,
            main_compile_receipt=compile_receipt,
            figure_summary=figures,
        )

    supplement, main_manuscript, _metadata_path, compile_receipt, figures = _final_fixture(
        tmp_path / "drift"
    )
    summary = json.loads(figures.read_text(encoding="utf-8"))
    first_record = next(iter(summary["supplementary_figures"].values()))
    first = Path(first_record["bundle"]["files"]["png"])
    first.write_bytes(first.read_bytes() + b"drift")
    with pytest.raises(builder.SupplementaryDocxError, match="PNG hash mismatch"):
        builder.build_supplementary_docx(
            mode="final",
            supplement=supplement,
            output=tmp_path / "blocked-drift.docx",
            main_manuscript=main_manuscript,
            main_compile_receipt=compile_receipt,
            figure_summary=figures,
        )


def test_anonymous_supplement_has_no_metadata_input_and_rolls_back(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    assert "submission_metadata" not in inspect.signature(
        builder.build_supplementary_docx
    ).parameters
    supplement, main_manuscript, _metadata_path, compile_receipt, figures = _final_fixture(
        tmp_path
    )
    output = tmp_path / "anonymous-supplement.docx"
    receipt = tmp_path / "anonymous-supplement.receipt.json"
    publish = builder.main._publish_no_overwrite
    calls = 0

    def fail_receipt(source: Path, destination: Path) -> None:
        nonlocal calls
        calls += 1
        if calls == 2:
            raise OSError("injected receipt publication failure")
        publish(source, destination)

    monkeypatch.setattr(builder.main, "_publish_no_overwrite", fail_receipt)
    with pytest.raises(OSError, match="injected"):
        builder.build_supplementary_docx(
            mode="final",
            supplement=supplement,
            output=output,
            main_manuscript=main_manuscript,
            main_compile_receipt=compile_receipt,
            figure_summary=figures,
            receipt=receipt,
        )
    assert not output.exists() and not receipt.exists()
